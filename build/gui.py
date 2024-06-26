
import subprocess

# List of required packages
required_packages = [
    "pymongo",
    "Pillow",
    "vidstream",
]

# Check if each package is installed, and install it if not
for package in required_packages:
    try:
        __import__(package)
    except ImportError:
        print(f"Installing {package}...")
        subprocess.check_call(["pip", "install", package])
from pathlib import Path
from tkinter import ttk
from tkinter import *
#from ttkbootstrap import Style
import pymongo
import base64
from tkinter import filedialog
from PIL import Image, ImageTk,ImageDraw
import urllib.parse
from tkinter import messagebox
# Explicit imports to satisfy Flake8
from tkinter import filedialog,Tk, Canvas, Entry, Text, Button, PhotoImage
from pymongo.mongo_client import MongoClient
from pymongo.server_api import ServerApi
from vidstream import *
import tkinter as tk
import socket
import threading
import customtkinter
import subprocess
from bson import Binary

try:
    import nltk
    import sklearn
    import pandas
    import docx
except ModuleNotFoundError:
    print("NLTK is not installed. Installing...")
    subprocess.check_call(["pip", "install", "nltk"])
    subprocess.check_call(["pip","install","scikit-learn"])
    subprocess.check_call(["pip","install","pandas"])
    subprocess.check_call(["pip","install","python-docx"])
    import nltk
    import sklearn
    import pandas
    import docx
import tkinter as tk
import nltk
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.svm import SVC
import pandas as pd
import pickle
from docx import Document
import os
from datetime import datetime
from tkcalendar import Calendar  # Install tkcalendar via pip: pip install tkcalendar
from bson.objectid import ObjectId
import logging
from tkcalendar import Calendar, DateEntry
import datetime
import time
from datetime import datetime




 # Install tkcalendar via pip: pip install tkcalendar








# Get the directory path of the current Python script
script_dir = os.path.dirname(os.path.realpath(__file__))
sub_dir = "assets\\frame0"
full_path = os.path.join(script_dir,sub_dir)
# Now 'script_dir' contains the directory path of the current Python script

#Connection String to connect to the MongoDb Atlas
conn_str = "mongodb+srv://root:812003@cluster0.fshfquh.mongodb.net/?retryWrites=true&w=majority"
client = MongoClient(conn_str, server_api=ServerApi('1'))

try:
    client.admin.command('ping')
    print("Pinged Your Deplyment. You successfully connected to MongoDB!")
except Exception as e:
    print(e)
AdminCollection = None
AdminDatabase = client["AdminDatabase"]
EmployeeDatabase = client["EmployeeDatabase"]
GroupChatDatabase = client["GroupChatDatabase"]
TodoListDatabase = client["TodoListDatabase"]
LeaveManagementDatabase = client["LeaveManagementDatabase"]
AppointmentSchedulingDatabase = client['AppointmentSchedulingDatabase']
DocumentDatabase = client['DocumentDatabase']


AdminCollection = AdminDatabase["ManagerCollection"]
EmployeeCollection = EmployeeDatabase["EmployeeCollection"]
EmployeeTaskCollection = EmployeeDatabase["EmployeeTaskCollection"]
GroupChatCollection = GroupChatDatabase["GroupChatCollection"]
TodoListCollection = TodoListDatabase["TodoListCollection"]
LeaveManagementCollection = LeaveManagementDatabase["LeaveManagementCollection"]
AppointmentSchedullingCollection = AppointmentSchedulingDatabase['AppointmentSchedulingCollection']
DocumentCollection = DocumentDatabase['DocumentCollection']

#print(client.list_database_names)

OUTPUT_PATH = Path(__file__).parent
ASSETS_PATH = OUTPUT_PATH / Path(full_path)

def relative_to_assets(path: str) -> Path:
    return ASSETS_PATH / Path(path)

window = Tk()
window.geometry("700x450")
window.configure(bg = "#FFFFFF")
window.iconbitmap("Skive_Logo.ico")
window.title("Skive")

invited_users = []
invited_users_ipaddress = []

def start_servers():
    try:
        # Get the absolute paths to the server.py and NotificationServer.py files
        server_path = os.path.abspath(os.path.join(os.getcwd(), "build", "server.py"))
        notification_server_path = os.path.abspath(os.path.join(os.getcwd(), "build", "NotificationServer.py"))
        
        # Start server.py on port 5555
        subprocess.Popen(["python", server_path, "--port", "12345"])
        
        # Start NotificationServer.py on a different port, e.g., 5556
        subprocess.Popen(["python", notification_server_path, "--port", "5555"])
        
        print("Servers started successfully!")
    except Exception as e:
        print("Error starting servers:", e)

def browse_photo():
    filename = filedialog.askopenfilename(title="Select Photo", filetypes=(("Image files", "*.jpg; *.jpeg; *.png"), ("All files", "*.*")))
    print(filename)
    return filename

def Managerbrowse_photo():
    filename = filedialog.askopenfilename(title="Select Photo", filetypes=(("Image files", "*.jpg; *.jpeg; *.png"), ("All files", "*.*")))
    print(filename)
    return filename

def upload_to_database(details):
    print(details)
    LeaveManagementCollection.insert_one(details)

def open_leave_application_window(CompanyName,UserID):
    global SubmitFormLeaveForm_image

    new_window = Toplevel(window)  # Create a new window
    new_window.title("Leave Application Form")  # Set the title for the new window
    new_window.geometry("510x450")
    def on_button_click():
        # Get the details from the widgets
        details = {
            "CompanyName": CompanyName,
            "UserName": UserID,
            "start_date": start_date_entry.get_date().strftime('%Y-%m-%d'),
            "reason_for_leave": entry_1.get(),
            "contact_information": entry_2.get(),
            "end_date": end_date_entry.get_date().strftime('%Y-%m-%d'),
            "Status": "Pending"
            # Add other details as needed
        }
        
        upload_to_database(details)
        new_window.destroy()

        
    canvas = Canvas(
        new_window,
        bg = "#173054",
        height = 450,
        width = 510,
        bd = 0,
        highlightthickness = 0,
        relief = "ridge"
    )

    canvas.place(x = 0, y = 0)
    canvas.create_rectangle(
        47.0,
        63.0,
        230.0,
        89.0,
        fill="#173054",
        outline="")

    entry_image_1 = PhotoImage(
        file=relative_to_assets("entry_56.png"))
    entry_bg_1 = canvas.create_image(
        138.5,
        280.0,
        image=entry_image_1
    )
    entry_1 = Entry(
        new_window,
        bd=0,
        bg="#ECECD9",
        fg="#000716",
        highlightthickness=0
    )
    entry_1.place(
        x=52.0,
        y=267.0,
        width=173.0,
        height=24.0
    )

    entry_image_2 = PhotoImage(
        file=relative_to_assets("entry_57.png"))
    entry_bg_2 = canvas.create_image(
        372.5,
        280.0,
        image=entry_image_2
    )
    entry_2 = Entry(
        new_window,
        bd=0,
        bg="#ECECD9",
        fg="#000716",
        highlightthickness=0
    )
    entry_2.place(
        x=286.0,
        y=267.0,
        width=173.0,
        height=24.0
    )

    canvas.create_rectangle(
        294.0,
        63.0,
        477.0,
        89.0,
        fill="#173054",
        outline="")

    canvas.create_text(
        49.0,
        41.0,
        anchor="nw",
        text="Start date:-",
        fill="#ECECD9",
        font=("LibreCaslonText Regular", 15 * -1)
    )

    start_date_entry = DateEntry(
        new_window,
        bd=0,
        bg="#ECECD9",
        fg="#000716",
        highlightthickness=0
    )
    start_date_entry.place(
        x=52.0,
        y=95.0,
        width=173.0,
        height=24.0
    )

    canvas.create_text(
        49.0,
        245.0,
        anchor="nw",
        text="Reason for leave",
        fill="#ECECD9",
        font=("LibreCaslonText Regular", 15 * -1)
    )

    canvas.create_text(
        286.0,
        245.0,
        anchor="nw",
        text="Contact information",
        fill="#ECECD9",
        font=("LibreCaslonText Regular", 15 * -1)
    )

    canvas.create_text(
        297.0,
        45.0,
        anchor="nw",
        text="End date:",
        fill="#ECECD9",
        font=("LibreCaslonText Regular", 15 * -1)
    )

    end_date_entry = DateEntry(
        new_window,
        bd=0,
        bg="#ECECD9",
        fg="#000716",
        highlightthickness=0
    )
    end_date_entry.place(
        x=286.0,
        y=95.0,
        width=173.0,
        height=24.0
    )

    canvas.create_text(
        21.0,
        17.0,
        anchor="nw",
        text="Leave application form",
        fill="#ECECD9",
        font=("LibreCaslonText Regular", 16 * -1)
    )

    SubmitFormLeaveForm_image = PhotoImage(
        file=relative_to_assets("button_48.png"))
    button_1 = Button(
        new_window,
        image=SubmitFormLeaveForm_image,
        borderwidth=0,
        highlightthickness=0,
        command=lambda: on_button_click(),
        relief="flat"
    )
    button_1.place(
        x=195.0,
        y=345.0,
        width=119.0,
        height=35.0
    )

def display_leave_details(canvas, leave_details):
    y_offset = 100  # Initial vertical offset for displaying details

    for leave in leave_details:
        # Displaying the details in the rectangle canvas
        canvas.create_text(50, y_offset, anchor="nw", text=f"Start Date: {leave['start_date']}")
        canvas.create_text(50, y_offset + 20, anchor="nw", text=f"End Date: {leave['end_date']}")
        canvas.create_text(400, y_offset, anchor="ne", text=f"Status: {leave['Status']}")

        # Incrementing the vertical offset for the next set of details
        y_offset += 50

def LeaveManagement(CompanyName,UserID):
    global LeaveManagementApplybutton_image_1
    LeaveManagementFrame = Frame(window,height=450, width=510)
    LeaveManagementFrame.place(x=36, y=0)
    canvas = Canvas(
    LeaveManagementFrame,
    bg = "#3A868F",
    height = 450,
    width = 510,
    bd = 0,
    highlightthickness = 0,
    relief = "ridge"
    )

    canvas.place(x = 0, y = 0)

    def update_leave_details():
        # Retrieve leave details from the database
        leave_details = LeaveManagementCollection.find({"UserName": UserID})  # Implement this function to retrieve details from the database
        # Clear the canvas
        canvas.delete("all")
        # Display leave details in the canvas
        display_leave_details(canvas, leave_details)
        # Schedule the next update after a certain interval (e.g., every 60 seconds)
        threading.Timer(30, update_leave_details).start()

    # Initial update of leave details
    update_leave_details()
    LeaveManagementApplybutton_image_1 = PhotoImage(
        file=relative_to_assets("button_47.png"))
    button_1 = Button(LeaveManagementFrame,
        image=LeaveManagementApplybutton_image_1,
        borderwidth=0,
        highlightthickness=0,
        command=lambda: open_leave_application_window(CompanyName,UserID),
        relief="flat"
    )
    button_1.place(
        x=399.0,
        y=18.0,
        width=94.0,
        height=35.0
    )
    # step1 on the main from of the user if they applied for any leave and whether it is approved or not
    # step2 display a button on the top where they want to apply for a new leave 
    # step3 if button click on then display a new window where user enters all the details related to the leave as soon as he clicks on the submmit button
    # step4 then notify the admin about the leave admin can choose to approve or disapprove or reply later or may be just notify the user and display a list of leave to the admin in leave management Portal
    # step5 if the admin approves then on the user side change the pending status to approved or decline based on the admin reply
    # leave dates with exact duration reason for leave contact infoction during leave document if required signature / authentication pass

def SubmitForm(CompanyName,username,Password,EmployeeFirstNameentry,EmployeeMiddleNameentry,EmployeeLastNameentry,EmployeeAgeentry,EmployeeGenderentry,EmployeeMobileNumberentry,EmployeeBirthDateentry,EmployeeDesignationentry,EmployeeEducationentry,EmployeeMailIDentry,EmployeeWorkExperienceentry,EmployeeSalaryentry,selected_listbox):
    selected_items = []
    for index in range(selected_listbox.size()):
        selected_items.append(selected_listbox.get(index))
    # Concatenate the items into a single string separated by commas
    selected_items_str = ', '.join(selected_items)
    
    First_Name = EmployeeFirstNameentry.get()
    Middle_Name = EmployeeMiddleNameentry.get()
    Last_Name = EmployeeLastNameentry.get()
    Age = EmployeeAgeentry.get()
    Gender = EmployeeGenderentry.get()
    Mobile_Number = EmployeeMobileNumberentry.get()
    Birthdate = EmployeeBirthDateentry.get()
    Designation = EmployeeDesignationentry.get()
    MailID = EmployeeMailIDentry.get()
    Highest_Education = EmployeeEducationentry.get()
    Work_Experience = EmployeeWorkExperienceentry.get()
    Previous_Salary = EmployeeSalaryentry.get()
    Skills =selected_items_str
    photo_path = browse_photo()
    
    photo_path = photo_path  # Get the path of the uploaded photo
    if CompanyName and username and Password and First_Name and Middle_Name and Last_Name and Age and Gender and Mobile_Number and Birthdate and Designation and MailID and Highest_Education and Work_Experience and Previous_Salary and photo_path and Skills:
        with open(photo_path, "rb") as photo_file:
            photo_data = base64.b64encode(photo_file.read()).decode('utf-8')
            data = {
                "CompanyName" : CompanyName,
                "UserName" : username,
                "Password" : Password,
                "First_Name" : First_Name,
                "Middle_Name" : Middle_Name,
                "Last_Name" : Last_Name,
                "Age": Age,
                "Gender" : Gender,
                "Mobile_Number" : Mobile_Number,
                "Birthdate" : Birthdate,
                "Designation" : Designation,
                "MailID" : MailID,
                "Highest_Education" : Highest_Education,
                "Work_Experience" : Work_Experience,
                "Previous_Salary" : Previous_Salary,
                "Skills" : Skills,
                "photo": photo_data  # Store the photo data directly in the document
            }
            EmployeeCollection.insert_one(data)
            print("Data saved successfully!")
            print(username)
            UserProfile(CompanyName,username)

def ChooseSkills(CompanyName,username,Password,EmployeeFirstNameentry,EmployeeMiddleNameentry,EmployeeLastNameentry,EmployeeAgeentry,EmployeeGenderentry,EmployeeMobileNumberentry,EmployeeBirthDateentry,EmployeeDesignationentry,EmployeeEducationentry,EmployeeMailIDentry,EmployeeWorkExperienceentry,EmployeeSalaryentry):
    global image_image_1,button_image_1,entry_image_1
    def on_configure(event):
        # Update scroll region to encompass the full canvas
        SelectedOptionCanvas.configure(scrollregion=SelectedOptionCanvas.bbox("all"))

    def on_drag_start(event):
        widget = event.widget
        item = widget.find_closest(event.x, event.y)
        item_text = widget.itemcget(item, "text")
        widget.drag_data = {"item": item, "x": event.x, "y": event.y, "text": item_text}

    def on_drag_motion(event):
        widget = event.widget
        x, y = event.x, event.y
        widget.move(widget.drag_data["item"], x - widget.drag_data["x"], y - widget.drag_data["y"])
        widget.drag_data["x"] = x
        widget.drag_data["y"] = y

    def on_drag_end(event):
        widget = event.widget
        item = widget.find_closest(event.x, event.y)
        item_text = widget.drag_data["text"]
        selected_listbox.insert(tk.END, item_text)
        del widget.drag_data

    def update_list(event):
        category = category_combobox.get()
        update_canvas(get_items_by_category(category))

    def update_canvas(items):
        SelectedOptionCanvas.delete("all")
        y_offset = 10
        for item in items:
            SelectedOptionCanvas.create_text(10, y_offset, anchor="w", text=item, font=("Helvetica", 12))
            y_offset += 25
        SelectedOptionCanvas.config(scrollregion=SelectedOptionCanvas.bbox("all"))

    def get_items_by_category(category):
        items_dict = {
            "Programming Languages": programming_languages,
            "Frameworks and Libraries": frameworks_and_libraries,
            "Mobile Development": mobile_development,
            "Web Development": web_development,
            "UI/UX Design": ui_ux_design,
            "Game Development": game_development,
            "Software Engineering": software_engineering,
            "Cloud Computing": cloud_computing,
            "Security": security,
            "Data Science and Analytics": data_science_and_analytics,
            "Database Management": database_management,
            "Development Tools": development_tools,
            "Blockchain": blockchain,
            "Internet of Things (IoT)": Iot,
            "AI and Robotics": AI_and_Robotics,
            "Virtual Reality and Augmented Reality (VR/AR)": VR_AR,
            "Graphics Design": graphics_design,
            "User Interface": user_interface,
            "Project Management": project_management,
            "Business Skills": business_skills,
            "Regulatory and Compliance": regulatory_compliance,
            "Sustainability": sustainability,
            "Networking": networking,
            "Systems": systems,
            "Real Estate": realestate,
            "Communication": communication,
            "Management and Leadership": management_and_leadership,
            "Personal Development": personal_development,
            "Legal and Ethical": legal_and_ethical,
            "Technical Proficiency": Technical_proficiency,
            "Digital Collaboration": Digital_collaboration
        }
        return items_dict.get(category, [])

    # Define available items
    programming_languages = ["C#", "Java", "JavaScript", "Python", "PHP", "Kotlin", "Swift", "C++", "Bash", "R", "SQL"]
    frameworks_and_libraries = ["NET", "Unity", "TensorFlow", "Keras", "React Native", "Xamarin", "Angular", "Node.js", "React", "Flask", "D3.js", "OpenCV"]
    mobile_development = ["Android Development", "iOS Development (Swift, Kotlin)", "React Native", "Xamarin", "Mobile Development Best Practices", "Flutter", "HealthKit", "CareKit"]
    web_development = ["HTML", "CSS", "JavaScript", "SEO Tools", "Web Development", "Web Scraping", "Web Crawling", "HTML5", "CSS3", "Drupal", "WordPress"]
    ui_ux_design = ["Sketch", "Figma", "Adobe XD", "User Interface Design", "User Experience Design", "UI/UX Design"]
    game_development = ["Unity", "Game Design Principles"]
    software_engineering = ["Software Maintenance", "Enterprise Architecture", "Algorithm Design", "Application Development", "System Integration", "Application Design"]
    cloud_computing = ["AWS", "Azure", "Cloud Computing Fundamentals", "Serverless Architecture", "Cloud Storage Solutions"]
    security = ["Cybersecurity", "Encryption", "SSL", "Payment Processing APIs", "Network Security", "Ethical Hacking", "Data Security", "Information Security", "Authentication", "Biometric Identification"]
    data_science_and_analytics = ["Machine Learning", "Data Analysis", "Big Data Analytics", "Data Analytics", "Data Visualization", "Data Science", "Statistical Analysis", "Predictive Modeling", "Fraud Detection Algorithms", "Health Informatics", "Quantitative Analysis", "Financial Modeling", "Statistics"]
    database_management = ["Database Management", "MySQL", "MongoDB", "Oracle", "SQL Optimization", "Database Management", "Performance Tuning"]
    development_tools = ["Docker", "Kubernetes", "Selenium", "Test Automation Frameworks", "Version Control", "CI/CD Pipelines"]
    blockchain =["Solidity", "Ethereum", "Smart Contracts", "Blockchain Concepts", "DApp Development"]
    Iot =["IoT Platforms", "IoT Integration", "Sensor Networks", "Arduino", "Raspberry Pi"]
    AI_and_Robotics = ["AI Principles", "Machine Learning Algorithms", "Robotics", "Embedded Systems", "AI Concepts", "Computer Vision", "Image Processing", "Speech Recognition Libraries", "NLP Techniques", "Chatbot Frameworks"]
    VR_AR = ["VR SDKs"]
    graphics_design = ["Adobe Suite", "Blender", "Maya", "3D Graphics"]
    user_interface =["Sketch", "Figma", "Adobe XD"]
    project_management = ["Project Management Methodologies", "Agile", "Scrum", "Analytical Skills", "Team Collaboration", "Leadership"]
    business_skills = ["Market Research", "Customer Behavior Analysis", "Market Analysis", "Strategic Planning", "Business Strategy", "Operational Efficiency", "Financial Planning", "Fundraising Strategies", "Competitive Intelligence", "Strategic Analysis", "Product Development", "Creativity"]
    regulatory_compliance = ["Legal Knowledge", "Ethical Decision Making", "Regulatory Compliance", "Environmental Law", "Data Protection Policies", "Banking and Finance Knowledge"]
    sustainability = ["Sustainability Planning", "Environmental Science", "Sustainability Strategies"]
    networking = ["Networking", "Socket Programming", "Multi-thread Programming"]
    systems = ["Satellite Systems", "Communication Theory", "Signal Processing", "Circuit Analysis", "Circuit Design", "Soldering", "Energy Systems", "Mechanical Analysis"]
    realestate = ["Real Estate Knowledge"]
    communication =["Interpersonal Communication", "Active Listening", "Customer Service", "Empathy", "Communication Skills"]
    management_and_leadership = ["Organizational Development", "Leadership", "Employee Motivation Techniques", "Feedback Mechanisms", "Constructive Criticism", "Team Management", "Change Management"]
    personal_development = ["Time Management", "Stress Management", "Self-Discipline", "Self-Motivation Techniques", "Continuous Learning", "Adaptability"]
    legal_and_ethical = ["Legal Knowledge", "Ethical Decision Making"]
    Technical_proficiency = ["Technical Skills", "Technical Adaptability"]
    Digital_collaboration = ["Digital Collaboration Tools"]
    # Define other categories similarly

    # Create category dropdown
    categories = ["Programming Languages","Frameworks and libraries","Mobile Development","Web Development","UI/UX Design","Game Development","Software Engineering","Cloud Computing","Security","Data Science and Analytics","Database Management","Development Tools","Blockchain","IoT","AI and Robotics","VR/AR","Graphic Design","User Interface Design","Project Management","Business Skills","Regulatory and Compliance"," Sustainability","Networking","Systems","Real Estate","Communication","Management and Leadership","Personal Development","Legal and Ethical","Technical Proficiency","Digital Collaboration"]



    canvas = tk.Canvas(
        window,
        bg="#173054",
        height=450,
        width=700,
        bd=0,
        highlightthickness=0,
        relief="ridge"
    )
    canvas.place(x=0, y=0)

    image_image_1 = tk.PhotoImage(file=relative_to_assets("image_16.png"))
    image_1 = canvas.create_image(350.0, 225.0, image=image_image_1)
    
    canvas.create_text(
    162.0,
    34.0,
    anchor="nw",
    text="Choose Your Skill :",
    fill="#FFFFFF",
    font=("LibreCaslonText Regular", 17 * -1)
)

    entry_image_1 = tk.PhotoImage(file=relative_to_assets("entry_54.png"))
    entry_bg_1 = canvas.create_image(472.5, 44.5, image=entry_image_1)

    category_combobox = ttk.Combobox(canvas, values=categories)
    category_combobox.place(x=357.0, y=29.0, width=231.0, height=29.0)
    category_combobox.bind("<<ComboboxSelected>>", update_list)

    canvas.create_text(
        47.0,
        85.0,
        anchor="nw",
        text="Available options:-",
        fill="#FFFFFF",
        font=("LibreCaslonText Regular", 13 * -1)
    )

    canvas.create_text(
        392.0,
        84.0,
        anchor="nw",
        text="Selected options:-",
        fill="#FFFFFF",
        font=("LibreCaslonText Regular", 13 * -1)
    )

    canvas_frame = ttk.Frame(canvas, width=303.0, height=253.0)
    canvas_frame.place(x=26.0, y=103.0)

    SelectedOptionCanvas = tk.Canvas(canvas_frame, bg="#3A868F")
    SelectedOptionCanvas.place(x=0, y=0, height=247, width=296)

    vsb = tk.Scrollbar(SelectedOptionCanvas, orient="vertical", command=SelectedOptionCanvas.yview)
    vsb.pack(side="right", fill="y")
    SelectedOptionCanvas.config(yscrollcommand=vsb.set)
    selected_listbox = tk.Listbox(canvas, width=36, height=11, selectmode=tk.MULTIPLE)
    selected_listbox.place(x=377, y=106)
    vsb_listbox = tk.Scrollbar(window, orient="vertical", command=selected_listbox.yview)
    vsb_listbox.place(x=selected_listbox.winfo_x() + selected_listbox.winfo_width(), y=selected_listbox.winfo_y(), height=selected_listbox.winfo_height())
    selected_listbox.config(yscrollcommand=vsb_listbox.set)


    # Bind the events directly to the SelectedOptionCanvas
    SelectedOptionCanvas.bind("<ButtonPress-1>", on_drag_start)
    SelectedOptionCanvas.bind("<B1-Motion>", on_drag_motion)
    SelectedOptionCanvas.bind("<ButtonRelease-1>", on_drag_end)

    button_image_1 = PhotoImage(
    file=relative_to_assets("button_35.png"))
    button_1 = Button(
        image=button_image_1,
        borderwidth=0,
        highlightthickness=0,
        command=lambda: SubmitForm(CompanyName,username,Password,EmployeeFirstNameentry,EmployeeMiddleNameentry,EmployeeLastNameentry,EmployeeAgeentry,EmployeeGenderentry,EmployeeMobileNumberentry,EmployeeBirthDateentry,EmployeeDesignationentry,EmployeeEducationentry,EmployeeMailIDentry,EmployeeWorkExperienceentry,EmployeeSalaryentry,selected_listbox),
        relief="flat"
    )
    button_1.place(
        x=287.0,
        y=385.0,
        width=126.0,
        height=36.0
    )

def display(selected_listbox):
    # Create an empty list to store the items
    selected_items = []

    # Iterate over each item in the selected_listbox and append it to the selected_items list
    for index in range(selected_listbox.size()):
        selected_items.append(selected_listbox.get(index))

    # Print the list of selected items
    print(selected_items)

def PersonalDetailForm(CompanyName,username,Password):
    global image_image_1,entry_image_1,entry_image_2,entry_image_3,entry_image_4,entry_image_5,entry_image_6,entry_image_7,entry_image_8,entry_image_9,entry_image_10,entry_image_11,entry_image_12,button_image_1,button_image_2
    canvas = Canvas(
    window,
    bg = "#225777",
    height = 450,
    width = 700,
    bd = 0,
    highlightthickness = 0,
    relief = "ridge"
    )

    canvas.place(x = 0, y = 0)
    image_image_1 = PhotoImage(
        file=relative_to_assets("image_15.png"))
    image_1 = canvas.create_image(
        350.0,
        225.0,
        image=image_image_1
    )

    canvas.create_text(
        72.0,
        56.0,
        anchor="nw",
        text="Personal Details",
        fill="#173054",
        font=("LibreCaslonText Regular", 14 * -1)
    )

    canvas.create_text(
        72.0,
        225.0,
        anchor="nw",
        text="Work Detials",
        fill="#173054",
        font=("LibreCaslonText Regular", 14 * -1)
    )

    canvas.create_text(
        88.0,
        86.0,
        anchor="nw",
        text="First Name",
        fill="#173054",
        font=("LibreCaslonText Regular", 11 * -1)
    )

    canvas.create_text(
        89.0,
        137.0,
        anchor="nw",
        text="Age",
        fill="#173054",
        font=("LibreCaslonText Regular", 11 * -1)
    )

    canvas.create_text(
        89.0,
        187.0,
        anchor="nw",
        text="Birthdate(DD/MM/YYYY)",
        fill="#173054",
        font=("LibreCaslonText Regular", 11 * -1)
    )

    canvas.create_text(
        87.0,
        249.0,
        anchor="nw",
        text="Email ID",
        fill="#173054",
        font=("LibreCaslonText Regular", 11 * -1)
    )

    canvas.create_text(
        281.0,
        249.0,
        anchor="nw",
        text="Work Experience in Years",
        fill="#173054",
        font=("LibreCaslonText Regular", 11 * -1)
    )

    canvas.create_text(
        475.0,
        249.0,
        anchor="nw",
        text="Previous Salary",
        fill="#173054",
        font=("LibreCaslonText Regular", 11 * -1)
    )

    canvas.create_text(
        281.0,
        187.0,
        anchor="nw",
        text="Designation",
        fill="#173054",
        font=("LibreCaslonText Regular", 11 * -1)
    )

    canvas.create_text(
        281.0,
        137.0,
        anchor="nw",
        text="Gender",
        fill="#173054",
        font=("LibreCaslonText Regular", 11 * -1)
    )

    canvas.create_text(
        473.0,
        137.0,
        anchor="nw",
        text="Mobile Number",
        fill="#173054",
        font=("LibreCaslonText Regular", 11 * -1)
    )

    canvas.create_text(
        473.0,
        187.0,
        anchor="nw",
        text="Highest Education",
        fill="#173054",
        font=("LibreCaslonText Regular", 11 * -1)
    )

    canvas.create_text(
        281.0,
        87.0,
        anchor="nw",
        text="Middle Name",
        fill="#173054",
        font=("LibreCaslonText Regular", 11 * -1)
    )

    canvas.create_text(
        474.0,
        86.0,
        anchor="nw",
        text="Last Name",
        fill="#173054",
        font=("LibreCaslonText Regular", 11 * -1)
    )

    entry_image_1 = PhotoImage(
        file=relative_to_assets("entry_42.png"))
    entry_bg_1 = canvas.create_image(
        158.0,
        110.5,
        image=entry_image_1
    )
    EmployeeFirstNameentry = Entry(
        bd=0,
        bg="#173054",
        fg="#000716",
        highlightthickness=0
    )
    EmployeeFirstNameentry.place(
        x=90.0,
        y=101.0,
        width=136.0,
        height=17.0
    )

    entry_image_2 = PhotoImage(
        file=relative_to_assets("entry_43.png"))
    entry_bg_2 = canvas.create_image(
        350.0,
        110.5,
        image=entry_image_2
    )
    EmployeeMiddleNameentry = Entry(
        bd=0,
        bg="#173054",
        fg="#000716",
        highlightthickness=0
    )
    EmployeeMiddleNameentry.place(
        x=282.0,
        y=101.0,
        width=136.0,
        height=17.0
    )

    entry_image_3 = PhotoImage(
        file=relative_to_assets("entry_44.png"))
    entry_bg_3 = canvas.create_image(
        542.0,
        110.5,
        image=entry_image_3
    )
    EmployeeLastNameentry = Entry(
        bd=0,
        bg="#173054",
        fg="#000716",
        highlightthickness=0
    )
    EmployeeLastNameentry.place(
        x=474.0,
        y=101.0,
        width=136.0,
        height=17.0
    )

    entry_image_4 = PhotoImage(
        file=relative_to_assets("entry_45.png"))
    entry_bg_4 = canvas.create_image(
        158.0,
        160.5,
        image=entry_image_4
    )
    EmployeeAgeentry = Entry(
        bd=0,
        bg="#173054",
        fg="#000716",
        highlightthickness=0
    )
    EmployeeAgeentry.place(
        x=90.0,
        y=151.0,
        width=136.0,
        height=17.0
    )

    entry_image_5 = PhotoImage(
        file=relative_to_assets("entry_46.png"))
    entry_bg_5 = canvas.create_image(
        350.0,
        160.5,
        image=entry_image_5
    )
    EmployeeGenderentry = Entry(
        bd=0,
        bg="#173054",
        fg="#000716",
        highlightthickness=0
    )
    EmployeeGenderentry.place(
        x=282.0,
        y=151.0,
        width=136.0,
        height=17.0
    )

    entry_image_6 = PhotoImage(
        file=relative_to_assets("entry_47.png"))
    entry_bg_6 = canvas.create_image(
        542.0,
        160.5,
        image=entry_image_6
    )
    EmployeeMobileNumberentry = Entry(
        bd=0,
        bg="#173054",
        fg="#000716",
        highlightthickness=0
    )
    EmployeeMobileNumberentry.place(
        x=474.0,
        y=151.0,
        width=136.0,
        height=17.0
    )

    entry_image_7 = PhotoImage(
        file=relative_to_assets("entry_48.png"))
    entry_bg_7 = canvas.create_image(
        158.0,
        210.5,
        image=entry_image_7
    )
    EmployeeBirthDateentry = Entry(
        bd=0,
        bg="#173054",
        fg="#000716",
        highlightthickness=0
    )
    EmployeeBirthDateentry.place(
        x=90.0,
        y=201.0,
        width=136.0,
        height=17.0
    )

    entry_image_8 = PhotoImage(
        file=relative_to_assets("entry_49.png"))
    entry_bg_8 = canvas.create_image(
        350.0,
        210.5,
        image=entry_image_8
    )
    EmployeeDesignationentry = Entry(
        bd=0,
        bg="#173054",
        fg="#000716",
        highlightthickness=0
    )
    EmployeeDesignationentry.place(
        x=282.0,
        y=201.0,
        width=136.0,
        height=17.0
    )

    entry_image_9 = PhotoImage(
        file=relative_to_assets("entry_50.png"))
    entry_bg_9 = canvas.create_image(
        542.0,
        210.5,
        image=entry_image_9
    )
    EmployeeEducationentry = Entry(
        bd=0,
        bg="#173054",
        fg="#000716",
        highlightthickness=0
    )
    EmployeeEducationentry.place(
        x=474.0,
        y=201.0,
        width=136.0,
        height=17.0
    )

    entry_image_10 = PhotoImage(
        file=relative_to_assets("entry_51.png"))
    entry_bg_10 = canvas.create_image(
        158.0,
        272.5,
        image=entry_image_10
    )
    EmployeeMailIDentry = Entry(
        bd=0,
        bg="#173054",
        fg="#000716",
        highlightthickness=0
    )
    EmployeeMailIDentry.place(
        x=90.0,
        y=263.0,
        width=136.0,
        height=17.0
    )

    entry_image_11 = PhotoImage(
        file=relative_to_assets("entry_52.png"))
    entry_bg_11 = canvas.create_image(
        350.0,
        272.5,
        image=entry_image_11
    )
    EmployeeWorkExperienceentry = Entry(
        bd=0,
        bg="#173054",
        fg="#000716",
        highlightthickness=0
    )
    EmployeeWorkExperienceentry.place(
        x=282.0,
        y=263.0,
        width=136.0,
        height=17.0
    )

    entry_image_12 = PhotoImage(
        file=relative_to_assets("entry_53.png"))
    entry_bg_12 = canvas.create_image(
        542.0,
        272.5,
        image=entry_image_12
    )
    EmployeeSalaryentry = Entry(
        bd=0,
        bg="#173054",
        fg="#000716",
        highlightthickness=0
    )
    EmployeeSalaryentry.place(
        x=474.0,
        y=263.0,
        width=136.0,
        height=17.0
    )
    
    def call():
        filename = browse_photo()
        return filename

    button_image_1 = PhotoImage(
        file=relative_to_assets("button_33.png"))
    button_1 = Button(
        image=button_image_1,
        borderwidth=0,
        highlightthickness=0,
        command=lambda:call(),
        relief="flat"
    )
    button_1.place(
        x=113.0,
        y=330.0,
        width=118.0,
        height=28.0
    )
    

    button_image_2 = PhotoImage(
        file=relative_to_assets("button_34.png"))
    button_2 = Button(
        image=button_image_2,
        borderwidth=0,
        highlightthickness=0,
        command=lambda: ChooseSkills(CompanyName,username,Password,EmployeeFirstNameentry,EmployeeMiddleNameentry,EmployeeLastNameentry,EmployeeAgeentry,EmployeeGenderentry,EmployeeMobileNumberentry,EmployeeBirthDateentry,EmployeeDesignationentry,EmployeeEducationentry,EmployeeMailIDentry,EmployeeWorkExperienceentry,EmployeeSalaryentry),
        relief="flat"
    )
    button_2.place(
        x=469.0,
        y=330.0,
        width=118.0,
        height=28.0
    )
    


    '''print(username)
    print(Password)
    canvas = Canvas(
    window,
    bg = "#225777",
    height = 450,
    width = 700,
    bd = 0,
    highlightthickness = 0,
    relief = "ridge"
    )

    canvas.place(x = 0, y = 0)
    canvas.create_rectangle(
        57.0,
        41.0,
        642.0,
        408.0,
        fill="#D9D9D9",
        outline="")

    canvas.create_text(
        72.0,
        56.0,
        anchor="nw",
        text="Personal Details",
        fill="#173054",
        font=("Libre Caslon Text", 14 * -1)
    )

    canvas.create_text(
        72.0,
        243.0,
        anchor="nw",
        text="Work Details",
        fill="#173054",
        font=("Libre Caslon Text", 14 * -1)
    )

    canvas.create_text(
        88.0,
        86.0,
        anchor="nw",
        text="First Name",
        fill="#173054",
        font=("Libre Caslon Text", 11 * -1)
    )

    canvas.create_text(
        89.0,
        137.0,
        anchor="nw",
        text="Age",
        fill="#173054",
        font=("Libre Caslon Text", 11 * -1)
    )

    canvas.create_text(
        89.0,
        187.0,
        anchor="nw",
        text="Birthdate(DD/MM/YYYY)",
        fill="#173054",
        font=("Libre Caslon Text", 11 * -1)
    )

    canvas.create_text(
        89.0,
        266.0,
        anchor="nw",
        text="Work Experience",
        fill="#173054",
        font=("Libre Caslon Text", 11 * -1)
    )

    canvas.create_text(
        277.0,
        266.0,
        anchor="nw",
        text="Previous Salary",
        fill="#173054",
        font=("Libre Caslon Text", 11 * -1)
    )

    canvas.create_text(
        281.0,
        137.0,
        anchor="nw",
        text="Gender",
        fill="#173054",
        font=("Libre Caslon Text", 11 * -1)
    )

    canvas.create_text(
        281.0,
        187.0,
        anchor="nw",
        text="Address",
        fill="#173054",
        font=("Libre Caslon Text", 11 * -1)
    )

    canvas.create_text(
        473.0,
        137.0,
        anchor="nw",
        text="Mobile Number",
        fill="#173054",
        font=("Libre Caslon Text", 11 * -1)
    )

    canvas.create_text(
        473.0,
        187.0,
        anchor="nw",
        text="Highest Education",
        fill="#173054",
        font=("Libre Caslon Text", 11 * -1)
    )

    canvas.create_text(
        474.0,
        266.0,
        anchor="nw",
        text="Skills",
        fill="#173054",
        font=("Libre Caslon Text", 11 * -1)
    )

    canvas.create_text(
        281.0,
        87.0,
        anchor="nw",
        text="Middle Name",
        fill="#173054",
        font=("Libre Caslon Text", 11 * -1)
    )

    canvas.create_text(
        474.0,
        86.0,
        anchor="nw",
        text="Last Name",
        fill="#173054",
        font=("Libre Caslon Text", 11 * -1)
    )

    entry_image_1 = PhotoImage(
        file=relative_to_assets("entry_12.png"))
    entry_bg_1 = canvas.create_image(
        158.0,
        110.5,
        image=entry_image_1
    )
    FirstName_entry = Entry(
        bd=0,
        bg="#173054",
        fg="#ECECD9",
        highlightthickness=0
    )
    FirstName_entry.place(
        x=90.0,
        y=101.0,
        width=136.0,
        height=17.0
    )

    entry_image_2 = PhotoImage(
        file=relative_to_assets("entry_13.png"))
    entry_bg_2 = canvas.create_image(
        158.0,
        160.5,
        image=entry_image_2
    )
    Age_entry = Entry(
        bd=0,
        bg="#173054",
        fg="#ECECD9",
        highlightthickness=0
    )
    Age_entry.place(
        x=90.0,
        y=151.0,
        width=136.0,
        height=17.0
    )

    entry_image_3 = PhotoImage(
        file=relative_to_assets("entry_14.png"))
    entry_bg_3 = canvas.create_image(
        158.0,
        210.5,
        image=entry_image_3
    )
    Birthdate_entry = Entry(
        bd=0,
        bg="#173054",
        fg="#ECECD9",
        highlightthickness=0
    )
    Birthdate_entry.place(
        x=90.0,
        y=201.0,
        width=136.0,
        height=17.0
    )

    entry_image_4 = PhotoImage(
        file=relative_to_assets("entry_15.png"))
    entry_bg_4 = canvas.create_image(
        158.0,
        289.5,
        image=entry_image_4
    )
    WorkExperience_entry = Entry(
        bd=0,
        bg="#173054",
        fg="#ECECD9",
        highlightthickness=0
    )
    WorkExperience_entry.place(
        x=90.0,
        y=280.0,
        width=136.0,
        height=17.0
    )

    entry_image_5 = PhotoImage(
        file=relative_to_assets("entry_16.png"))
    entry_bg_5 = canvas.create_image(
        350.0,
        217.0,
        image=entry_image_5
    )
    Address_entry = Text(
        bd=0,
        bg="#173054",
        fg="#ECECD9",
        highlightthickness=0
    )
    Address_entry.place(
        x=282.0,
        y=201.0,
        width=136.0,
        height=30.0
    )

    entry_image_6 = PhotoImage(
        file=relative_to_assets("entry_17.png"))
    entry_bg_6 = canvas.create_image(
        350.0,
        289.5,
        image=entry_image_6
    )
    PreviousSalary_entry = Entry(
        bd=0,
        bg="#173054",
        fg="#ECECD9",
        highlightthickness=0
    )
    PreviousSalary_entry.place(
        x=282.0,
        y=280.0,
        width=136.0,
        height=17.0
    )

    entry_image_7 = PhotoImage(
        file=relative_to_assets("entry_18.png"))
    entry_bg_7 = canvas.create_image(
        350.0,
        160.5,
        image=entry_image_7
    )
    Gender_entry = Entry(
        bd=0,
        bg="#173054",
        fg="#ECECD9",
        highlightthickness=0
    )
    Gender_entry.place(
        x=282.0,
        y=151.0,
        width=136.0,
        height=17.0
    )

    entry_image_8 = PhotoImage(
        file=relative_to_assets("entry_19.png"))
    entry_bg_8 = canvas.create_image(
        542.0,
        160.5,
        image=entry_image_8
    )
    MobileNumber_entry = Entry(
        bd=0,
        bg="#173054",
        fg="#ECECD9",
        highlightthickness=0
    )
    MobileNumber_entry.place(
        x=474.0,
        y=151.0,
        width=136.0,
        height=17.0
    )

    entry_image_9 = PhotoImage(
        file=relative_to_assets("entry_20.png"))
    entry_bg_9 = canvas.create_image(
        542.0,
        210.5,
        image=entry_image_9
    )
    HighestEducation_entry = Entry(
        bd=0,
        bg="#173054",
        fg="#ECECD9",
        highlightthickness=0
    )
    HighestEducation_entry.place(
        x=474.0,
        y=201.0,
        width=136.0,
        height=17.0
    )

    entry_image_10 = PhotoImage(
        file=relative_to_assets("entry_21.png"))
    entry_bg_10 = canvas.create_image(
        542.0,
        289.5,
        image=entry_image_10
    )
    Skills_entry = Entry(
        bd=0,
        bg="#173054",
        fg="#ECECD9",
        highlightthickness=0
    )
    Skills_entry.place(
        x=474.0,
        y=280.0,
        width=136.0,
        height=17.0
    )

    entry_image_11 = PhotoImage(
        file=relative_to_assets("entry_22.png"))
    entry_bg_11 = canvas.create_image(
        350.0,
        110.5,
        image=entry_image_11
    )
    MiddleName_entry = Entry(
        bd=0,
        bg="#173054",
        fg="#ECECD9",
        highlightthickness=0
    )
    MiddleName_entry.place(
        x=282.0,
        y=101.0,
        width=136.0,
        height=17.0
    )

    entry_image_12 = PhotoImage(
        file=relative_to_assets("entry_23.png"))
    entry_bg_12 = canvas.create_image(
        542.0,
        110.5,
        image=entry_image_12
    )
    LastName_entry = Entry(
        bd=0,
        bg="#173054",
        fg="#ECECD9",
        highlightthickness=0
    )
    LastName_entry.place(
        x=474.0,
        y=101.0,
        width=136.0,
        height=17.0
    )

    button_image_1 = PhotoImage(
        file=relative_to_assets("button_25.png"))
    button_1 = Button(
        image=button_image_1,
        borderwidth=0,
        highlightthickness=0,
        command=lambda:SubmitForm(CompanyName,username,Password,FirstName_entry,Age_entry,Birthdate_entry,WorkExperience_entry,Address_entry,PreviousSalary_entry,Gender_entry,MobileNumber_entry,HighestEducation_entry,Skills_entry,MiddleName_entry,LastName_entry),
        
        relief="flat"
    )
    button_1.place(
        x=113.0,
        y=330.0,
        width=118.0,
        height=28.0
    )

    button_image_2 = PhotoImage(
        file=relative_to_assets("button_26.png"))
    button_2 = Button(
        image=button_image_2,
        borderwidth=0,
        highlightthickness=0,
        command=lambda :browse_photo(),
        
        relief="flat"
    )
    button_2.place(
        x=469.0,
        y=330.0,
        width=118.0,
        height=28.0
    )'''

def VideoConferencing():
    frame_width=357
    frame_height=450

    left_space = (window.winfo_width()- 192 - 151)/ 2

    frame = Frame(window, bg="#ECECD9", bd=0, relief=SOLID)
    frame.place(relx=left_space / window.winfo_width(),rely=0.5, anchor=W, width=frame_width, height=frame_height)
    local_ip_address = socket.gethostbyname(socket.gethostname())
    server = StreamingServer(local_ip_address, 9999)
    receiver = AudioReceiver(local_ip_address, 8888)

    def start_listening():
        t1= threading.Thread(target=server.start_server)
        t2= threading.Thread(target=receiver.start_server)
        t1.start()
        t2.start()


    def start_camera_stream():
        camera_client = CameraClient(text_target_ip.get(1.0,'end-1c'),7777)
        t3 = threading.Thread(target=camera_client.start_stream)
        t3.start()

    def start_screen_sharing():
        screen_client = ScreenShareClient(text_target_ip.get(1.0,'end-1c'),7777)
        t4 = threading.Thread(target=screen_client.start_stream)
        t4.start()

    def start_audio_stream():
        audio_sender = AudioSender(text_target_ip.get(1.0,'end-1c'),6666)
        t5 = threading.Thread(target=audio_sender.start_stream)
        t5.start()

    label_target_ip = tk.Label(frame, text="target IP:")
    label_target_ip.pack()

    text_target_ip = tk.Text(frame, height=1)
    text_target_ip.pack()

    btn_listen = tk.Button(frame,text="Start listening", width=50, command=start_listening)
    btn_listen.pack(anchor=tk.CENTER,expand=True)

    btn_camera = tk.Button(frame,text="Start Camera stream", width=50, command=start_camera_stream)
    btn_camera.pack(anchor=tk.CENTER,expand=True)

    btn_screen = tk.Button(frame,text="Start screen sharing", width=50, command=start_screen_sharing)
    btn_screen.pack(anchor=tk.CENTER,expand=True)

    btn_audio = tk.Button(frame,text="Start Audio Stream", width=50, command=start_audio_stream)
    btn_audio.pack(anchor=tk.CENTER,expand=True)
    window.after(1000, start_camera_stream)

def Todo(UserName, CompanyName):
    global DeleteTaskButton, AddTaskButton, TodoListBGImage
    TodoListMainFrame = Frame(window, height=450, width=510)
    TodoListMainFrame.place(x=36, y=0)

    canvas = Canvas(
        TodoListMainFrame,
        bg="#3A868F",
        height=450,
        width=510,
        bd=0,
        highlightthickness=0,
        relief="ridge"
    )

    canvas.place(x=0, y=0)
    
    def add_task_window():
        add_window = tk.Toplevel(window)
        add_window.title("Add Task")

        def submit_task():
            selected_date = cal.get_date()
            selected_hour = hour_var.get()
            selected_minute = minute_var.get()
            task = task_entry.get()

            if selected_date and selected_hour and selected_minute and task:
                datetime_str = f"{selected_date} {selected_hour}:{selected_minute}"
                datetime_obj = datetime.strptime(datetime_str, "%Y-%m-%d %H:%M")

                task_data = {
                    "username": UserName,
                    "company": CompanyName,
                    "date": datetime_obj,
                    "task": task
                }

                TodoListCollection.insert_one(task_data)
                print("Task added successfully.")

                add_window.destroy()
                display_tasks()
            else:
                print("Please fill in all fields.")

        try:
            cal = Calendar(add_window, selectmode="day", date_pattern="yyyy-mm-dd")
            cal.pack(pady=10)
        except AttributeError as e:
            print("Error: Calendar object has no attribute '_properties'.", e)

        hour_label = Label(add_window, text="Due Time (HH:MM):")
        hour_label.pack()
        
        hour_var = tk.StringVar(add_window)
        hour_dropdown = ttk.Combobox(add_window, textvariable=hour_var)
        hour_dropdown['values'] = tuple(str(i).zfill(2) for i in range(24))
        hour_dropdown.current(0)  # Default value
        hour_dropdown.pack()

        minute_var = tk.StringVar(add_window)
        minute_dropdown = ttk.Combobox(add_window, textvariable=minute_var)
        minute_dropdown['values'] = tuple(str(i).zfill(2) for i in range(60))
        minute_dropdown.current(0)  # Default value
        minute_dropdown.pack()

        task_label = Label(add_window, text="Task:")
        task_label.pack()
        task_entry = Entry(add_window)
        task_entry.pack()

        submit_button = Button(add_window, text="Submit", command=submit_task)
        submit_button.pack()

    def delete_task(task_id, text_id, rect_id):
        TodoListCollection.delete_one({"_id": task_id})
        canvas.delete(text_id)  # Delete the task text
        canvas.delete(rect_id)  # Delete the rectangle

    def display_tasks():
        # Clear the existing tasks
        canvas.delete("all")

        # Retrieve all tasks from the database
        tasks = TodoListCollection.find().sort("date", 1)

        # Group tasks by date
        tasks_by_date = {}
        for task in tasks:
            task_date = task["date"].strftime("%d %B %Y")
            if task_date not in tasks_by_date:
                tasks_by_date[task_date] = []
            tasks_by_date[task_date].append(task)

        y_position = 40  # Initial y position for the task details

        # Display tasks grouped by date
        for date, tasks_list in tasks_by_date.items():
            # Display the date using create_text
            canvas.create_text(20, y_position, text=date, fill="#173054", font=("Libre Caslon Text", 13), anchor="w")
            y_position += 20  # Increment y position for the task details

            # Display task details within separate rectangles
            for task in tasks_list:
                # Create a rectangle for each task
                rect_width = 445  # Adjust this width based on your canvas width
                rect_height = 30  # Adjust this height as needed
                canvas.create_rectangle(20, y_position, 20 + rect_width, y_position + rect_height, fill="#173054", outline="")

                # Display task details within the rectangle
                task_details = f"{task['task']}  {task['date'].strftime('%H:%M')}"
                canvas.create_text(30, y_position + 15, text=task_details, fill="#ECECD9", font=("Libre Caslon Text", 12), anchor="w")

                y_position += rect_height + 5  # Add space between rectangles

            y_position += 10  # Add padding after the last rectangle

        canvas.pack()




    display_tasks()

    TodoListBGImage = PhotoImage(file=relative_to_assets("image_21.png"))
    image_1 = canvas.create_image(
        255.0,
        225.0,
        image=TodoListBGImage
    )

    AddTaskButton = PhotoImage(file=relative_to_assets("button_45.png"))
    button_1 = Button(
        TodoListMainFrame,
        image=AddTaskButton,
        borderwidth=0,
        highlightthickness=0,
        command=lambda: add_task_window(),
        relief="flat"
    )
    button_1.place(
        x=404.0,
        y=11.0,
        width=91.0,
        height=31.0
    )

    # Display tasks immediately after the Todo function is executed
    display_tasks()



    # DeleteTaskButton = PhotoImage(
    #     file=relative_to_assets("button_46.png"))
    # button_2 = Button(
    #     TodoListMainFrame,
    #     image=DeleteTaskButton,
    #     borderwidth=0,
    #     highlightthickness=0,
    #     command=lambda: print("button_2 clicked"),
    #     relief="flat"
    # )
    # button_2.place(
    #     x=441.0,
    #     y=72.0,
    #     width=20.0,
    #     height=20.0
    # )

    # canvas.create_rectangle(
    #     42.0,
    #     68.0,
    #     467.0,
    #     97.0,
    #     fill="#173054",
    #     outline="")

    # canvas.create_text(
    #     42.0,
    #     46.0,
    #     anchor="nw",
    #     text="13th March 2024 Time",
    #     fill="#173054",
    #     font=("LibreCaslonText Regular", 13 * -1)
    # )

    # canvas.create_rectangle(
    #     51.0,
    #     76.0,
    #     63.0,
    #     88.0,
    #     fill="#000000",
    #     outline="")

    # canvas.create_text(
    #     77.0,
    #     68.0,
    #     anchor="nw",
    #     text="Display the Task here",
    #     fill="#ECECD9",
    #     font=("LibreCaslonText Regular", 13 * -1)
    # )

def TextEditor():
    import Text    

def TaskAssignmentTool(UserID,Company_Name):
    global TaskAssignmentFrame,UploadButtonImage,TaskAssignmentBG
    TaskAssignmentFrame = Frame(window, height=450, width=510)
    TaskAssignmentFrame.place(x=36, y=0)
    nltk.download('punkt')
    nltk.download('stopwords')
    nltk.download('wordnet')
    script_dir = os.path.dirname(os.path.realpath(__file__))
    sub_dir = "assets\\final_combined_dataset_skive.csv"
    full_path = os.path.join(script_dir,sub_dir)
    # Now 'script_dir' contains the directory path of the current Python script

    # Load and preprocess dataset
    dataset = pd.read_csv(full_path)

    # Function to handle NaN values and preprocess text
    def preprocess_text(text):
        if pd.isnull(text):
            return ""
        else:
            tokens = word_tokenize(text)
            tokens = [word for word in tokens if word.isalnum()]
            stop_words = set(stopwords.words('english'))
            tokens = [word for word in tokens if not word in stop_words]
            lemmatizer = WordNetLemmatizer()
            tokens = [lemmatizer.lemmatize(word.lower()) for word in tokens]
            return ' '.join(tokens)

    # Apply preprocessing to problem statements
    dataset['processed_statement'] = dataset['problem_statement'].apply(preprocess_text)

    # Fill NaN values in dataset
    dataset.fillna('', inplace=True)  # Replace NaN values with an empty string

    # Vectorize problem statements using TF-IDF
    tfidf_vectorizer = TfidfVectorizer(max_features=1000)
    X = tfidf_vectorizer.fit_transform(dataset['processed_statement'])
    y = dataset['required_skills']

    # Train Support Vector Machine (SVM) classifier
    svm_classifier = SVC(kernel='linear')
    svm_classifier.fit(X, y)

    # Save the trained vectorizer and classifier
    with open('vectorizer.pkl', 'wb') as f:
        pickle.dump(tfidf_vectorizer, f)

    with open('classifier.pkl', 'wb') as f:
        pickle.dump(svm_classifier, f)
    
    def displayEmployeeWithRequiredSkills(skills_text):
        DisplayEmployeeWithRequiredSkillsFrame = customtkinter.CTkScrollableFrame(TaskAssignmentFrame, width=290)
        DisplayEmployeeWithRequiredSkillsFrame.place(x=11, y=110)
        DisplayEmployeeWithRequiredSkillsFrame.configure(height=20)
        print(Company_Name)
        print(UserID)
        # Query the database
        cursor = EmployeeCollection.find({"CompanyName": Company_Name})
        matched_employees = []
        required_skills = set(skill.strip().lower() for skill in skills_text.split(','))
        print(required_skills)
        matched_skills_count = {}
        matched_skills = {}

        # Filter employees with required skills
        for employee in cursor:
            # Split the employee's skills into individual skills and convert them to lowercase
            employee_skills = [skill.strip().lower() for skill in employee["Skills"].split(',')]
            # Find the matched skills, ignoring case
            matched_skills_list = [skill for skill in employee_skills if skill in required_skills]
            if matched_skills_list:
                matched_skills_count[employee["_id"]] = len(matched_skills_list)
                matched_skills[employee["_id"]] = matched_skills_list
        # Sort employees based on the number of matched skills
        sorted_employees = sorted(matched_skills_count.items(), key=lambda x: x[1], reverse=True)

        
            
        for employee_id, matched_count in sorted_employees:
            if matched_count > 0:
                employee = EmployeeCollection.find_one({"_id": employee_id})
                print("Name:", employee["UserName"], matched_skills[employee_id])
                print("Number of Matched Skills:", matched_count)
                print("Matched Skills:", ", ".join(matched_skills[employee_id]))
                print("-------------------")
        
            if "photo" in employee:
                photo_data_base64 = employee["photo"]
                # Decode Base64 data
                photo_data_binary = base64.b64decode(photo_data_base64)
                with open("user_photo.jpg", "wb") as f:
                    f.write(photo_data_binary)
                image = Image.open("user_photo.jpg")
                image = image.resize((35, 35), Image.LANCZOS)
                image = round_corners(image, 50)
            
            else:
                print("Photo not found for the specified user.")
            user_frame = ttk.Frame(DisplayEmployeeWithRequiredSkillsFrame, padding=5, relief="raised")
            user_frame.pack(fill=tk.BOTH, expand=True)
            
            photo = ImageTk.PhotoImage(image)
            image_label = tk.Label(user_frame, image=photo)
            image_label.image = photo  # Keep a reference to avoid garbage collection
            image_label.pack(anchor="w")

            label = ttk.Label(user_frame, text=""+employee["First_Name"])
            label.pack(side=tk.LEFT)
            
            button = ttk.Button(user_frame, text="Assign", command=lambda:print("Assign"))
            button.pack(side=tk.RIGHT)
    

    # Function to extract text from Word document
    def extract_text_from_docx(file_path):
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)

    # Function to predict required skills
    def predict_required_skills(problem_statement):
        processed_statement = preprocess_text(problem_statement)
        vectorized_statement = tfidf_vectorizer.transform([processed_statement])
        predicted_skills = svm_classifier.predict(vectorized_statement)
        return predicted_skills

    # GUI application
    def upload_document():
        file_path = filedialog.askopenfilename(filetypes=[("Word documents", "*.docx"), ("All files", "*.*")])
        if file_path:
            try:
                document_text = extract_text_from_docx(file_path)
                predicted_skills = predict_required_skills(document_text)
                skills_text = ', '.join(predicted_skills)
                canvas.itemconfig(result_label,text="" + str(skills_text))
                print(predicted_skills)
                print("Problem Statement:")
                print(document_text)
                print("Required Skills =" +skills_text)
                displayEmployeeWithRequiredSkills(skills_text)

            except Exception as e:
                print("Error occurred:", e)


    
    canvas = Canvas(
        TaskAssignmentFrame,
        bg = "#3A868F",
        height = 450,
        width = 510,
        bd = 0,
        highlightthickness = 0,
        relief = "ridge"
    )

    canvas.place(x = 0, y = 0)
    TaskAssignmentBG = PhotoImage(
        file=relative_to_assets("image_17.png"))
    image_1 = canvas.create_image(
        255.0,
        225.0,
        image=TaskAssignmentBG
    )

    canvas.create_rectangle(
        17.0,
        11.0,
        493.0,
        165.0,
        fill="#225777",
        outline="")

    canvas.create_text(
        35.0,
        26.0,
        anchor="nw",
        text="Task Assignment Tool",
        fill="#ECECD9",
        font=("Libre Caslon Text", 19 * -1)
    )

    canvas.create_text(
        35.0,
        71.0,
        anchor="nw",
        text="Choose Document \nfrom Files:-",
        fill="#FFFFFF",
        font=("Libre Caslon Text", 14 * -1)
    )

    canvas.create_text(
        35.0,
        113.0,
        anchor="nw",
        text="Skills Required :-",
        fill="#FFFFFF",
        font=("Libre Caslon Text", 14 * -1)
    )

    result_label=canvas.create_text(
        159.0,
        113.0,
        anchor="nw",
        text="",
        fill="#FFFFFF",
        font=("Libre Caslon Text", 14 * -1)
    )

    UploadButtonImage = PhotoImage(
        file=relative_to_assets("button_36.png"))
    upload_button = Button(
        image=UploadButtonImage,
        borderwidth=0,
        highlightthickness=0,
        command=lambda: upload_document(),
        relief="flat"
    )
    upload_button.place(
        x=199.0,
        y=71.0,
        width=192.0,
        height=36.0
    )

    canvas.create_rectangle(
        17.0,
        181.0,
        493.0,
        433.0,
        fill="#D9D9D9",
        outline="")

    canvas.create_rectangle(
        33.0,
        197.0,
        479.0,
        278.0,
        fill="#173054",
        outline="")

    canvas.create_rectangle(
        33.0,
        284.0,
        479.0,
        365.0,
        fill="#173054",
        outline="")

def LeaveApplicationDisplay(CompanyName,UserID):
    LeaveApplicationFrame = Frame(window, height="450", width="510")
    LeaveApplicationFrame.place(x=36,y=0)
        
    canvas = Canvas(
        LeaveApplicationFrame,
        bg = "#173054",
        height = 450,
        width = 510,
        bd = 0,
        highlightthickness = 0,
        relief = "ridge"
    )

    canvas.place(x = 0, y = 0)
    canvas.create_text(
        102.0,
        7.0,
        anchor="nw",
        text="Leave Applications",
        fill="#FFFFFF",
        font=("Libre Caslon Text", 18 * -1)
    )
    
    # Clear existing widgets
    for widget in canvas.winfo_children():
        widget.destroy()

    # Fetch entries from the database
    entries = LeaveManagementCollection.find({"CompanyName": CompanyName})

    # Create and place widgets for each entry
    y_offset = 150  # Initial y-coordinate for the first entry
    for idx, entry in enumerate(entries):
        # Create a rectangle
        canvas.create_rectangle(
            21.0,
            y_offset - 57.0,
            489.0,
            y_offset + 1.0,
            fill="#3A868F",
            outline=""
        )

        # Create and place username text
        canvas.create_text(
            30.0,
            y_offset - 52.0,
            anchor="nw",
            text=f"{entry['UserName']}",
            fill="#FFFFFF",
            font=("LibreCaslonText Regular", 17 * -1)
        )

        # Create and place start date text
        canvas.create_text(
            30.0,
            y_offset - 25.0,
            anchor="nw",
            text=f"{entry['start_date']}",
            fill="#FFFFFF",
            font=("LibreCaslonText Regular", 15 * -1)
        )

        # Create and place end date text
        canvas.create_text(
            140.0,
            y_offset - 25.0,
            anchor="nw",
            text=f"{entry['end_date']}",
            fill="#FFFFFF",
            font=("LibreCaslonText Regular", 15 * -1)
        )

        # Create dropdown menu
        status_var = StringVar(LeaveApplicationFrame)
        status_var.set(entry['Status'])  # Set default value from database
        status_menu = OptionMenu(LeaveApplicationFrame, status_var, "Approve", "Refuse")
        status_menu.place(x=260.0, y=y_offset - 45.0)

        # Create submit button
        def on_submit(entry_id):
            print(entry_id)
            new_status = status_var.get()
            LeaveManagementCollection.update_one({"_id": entry_id}, {"$set": {"Status": new_status}})
            #update_widgets()  # Update widgets after database update

        submit_button = Button(
            LeaveApplicationFrame,
            text="Submit",
            command=lambda entry_id=entry['_id']: on_submit(entry_id),
            relief="flat"
        )
        submit_button.place(x=397.0, y=y_offset - 52.0, width=81.0, height=36.0)

        # Update y_offset for next entry
        y_offset += 100

def AppointmentDisplay(Company_Name,UserID):
    
    AppointmentDisplayFrame = Frame(window, height="450", width="510")
    AppointmentDisplayFrame.place(x=36, y=0)
    print("frame placed")

    canvas = Canvas(
        AppointmentDisplayFrame,
        bg="#173054",
        height=450,
        width=510,
        bd=0,
        highlightthickness=0,
        relief="ridge"
    )
    canvas.place(x=0,y=0)
    print("Canvas placed")
    def update_status(appointment_id, new_status):
        AppointmentSchedullingCollection.update_one({"_id": appointment_id}, {"$set": {"Status": new_status}})

    def display_appointments(CompanyName, UserID):
        print("Display appointment function")

        canvas.delete("all")
        
        appointments = AppointmentSchedullingCollection.find({"CompanyName": CompanyName})
        y_offset = 20
        for appointment in appointments:
            reason = appointment["reason"]
            date = appointment["date"]
            time = appointment["time"]
            status = appointment["Status"]

            canvas.create_text(
                20,
                y_offset,
                anchor="nw",
                text=f"Reason: {reason}",
                fill="#FFFFFF",
                font=("LibreCaslonText Regular", 12),
            )

            canvas.create_text(
                20,
                y_offset + 20,
                anchor="nw",
                text=f"Date: {date}",
                fill="#FFFFFF",
                font=("LibreCaslonText Regular", 12),
            )

            canvas.create_text(
                20,
                y_offset + 40,
                anchor="nw",
                text=f"Time: {time}",
                fill="#FFFFFF",
                font=("LibreCaslonText Regular", 12),
            )

            canvas.create_text(
                20,
                y_offset + 60,
                anchor="nw",
                text=f"Status: {status}",
                fill="#FFFFFF",
                font=("LibreCaslonText Regular", 12),
            )

            approve_button = Button(
                AppointmentDisplayFrame,
                text="Approve",
                command=lambda appointment_id=appointment["_id"]: update_status(appointment_id, "Approved"),
                bg="#4CAF50",
                fg="white",
                font=("LibreCaslonText Regular", 10),
                relief="flat"
            )
            approve_button.place(x=400, y=y_offset + 20)

            refuse_button = Button(
                AppointmentDisplayFrame,
                text="Refuse",
                command=lambda appointment_id=appointment["_id"]: update_status(appointment_id, "Refused"),
                bg="#F44336",
                fg="white",
                font=("LibreCaslonText Regular", 10),
                relief="flat"
            )
            refuse_button.place(x=450, y=y_offset + 20)

            y_offset += 100  # Adjust as needed

        # Schedule the next update after 15 seconds
        threading.Timer(15, display_appointments, args=(CompanyName, UserID)).start()
    display_appointments(Company_Name,UserID)

def AdminProfile(UserID,Company_Name):

    global image_image_1,image_image_2,button_image_1,button_image_2,button_image_3,button_image_4,button_image_5,button_image_6
    global button_image_7,button_image_8,button_image_9,button_image_10,button_image_11,button_image_12
    user_data = AdminCollection.find_one({"UserName": UserID})
    if user_data and "photo" in user_data:
        photo_data_base64 = user_data["photo"]
        # Decode Base64 data
        photo_data_binary = base64.b64decode(photo_data_base64)
        with open("user_photo.jpg", "wb") as f:
            f.write(photo_data_binary)
        image = Image.open("user_photo.jpg")
        image = image.resize((112, 112), Image.LANCZOS)
        image = round_corners(image, 50)
        
    else:
        print("Photo not found for the specified user.")
        return None
    if UserID:
        
        canvas = Canvas(
        window,
        bg = "#173054",
        height = 450,
        width = 700,
        bd = 0,
        highlightthickness = 0,
        relief = "ridge"
    )

        canvas.place(x = 0, y = 0)
        image_image_1 = PhotoImage(
            file=relative_to_assets("image_10.png"))
        image_1 = canvas.create_image(
            350.0,
            225.0,
            image=image_image_1
        )

        canvas.create_rectangle(
            546.0,
            0.0,
            700.0,
            450.0,
            fill="#225777",
            outline="")

        canvas.create_text(
            575.0,
            150.0,
            anchor="nw",
            text=""+ user_data["UserName"],
            fill="#FFFFFF",
            font=("Libre Caslon Text", 16 * -1)
        )

        canvas.create_text(
            557.0,
            229.0,
            anchor="nw",
            text="Tasks",
            fill="#FFFFFF",
            font=("Libre Caslon Text", 12 * -1)
        )

        canvas.create_text(
            575.0,
            171.0,
            anchor="nw",
            text=""+ user_data["Designation"],
            fill="#FFFFFF",
            font=("Libre Caslon Text", 8 * -1)
        )

        canvas.create_text(
            557.0,
            343.0,
            anchor="nw",
            text="Todo List",
            fill="#FFFFFF",
            font=("Libre Caslon Text", 12 * -1)
        )

        image_image_2 = PhotoImage(
            file=relative_to_assets("image_11.png"))
        image_2 = canvas.create_image(
            625.0,
            79.0,
            image=image_image_2
        )

        canvas.create_rectangle(
            0.0,
            0.0,
            36.0,
            450.0,
            fill="#225777",
            outline="")

        button_image_1 = PhotoImage(
            file=relative_to_assets("button_13.png"))
        VideoCallButton = Button(
            image=button_image_1,
            borderwidth=0,
            highlightthickness=0,
            command=lambda: print("button_13 clicked"),
            relief="flat"
        )
        VideoCallButton.place(
            x=2.0,
            y=3.0,
            width=32.0,
            height=32.0
        )

        button_image_2 = PhotoImage(
            file=relative_to_assets("button_14.png"))
        MailButton = Button(
            image=button_image_2,
            borderwidth=0,
            highlightthickness=0,
            command=lambda: MailApplication(),
            relief="flat"
        )
        MailButton.place(
            x=2.0,
            y=35.0,
            width=32.0,
            height=32.0
        )

        button_image_3 = PhotoImage(
            file=relative_to_assets("button_15.png"))
        button_3 = Button(
            image=button_image_3,
            borderwidth=0,
            highlightthickness=0,
            command=lambda:Display_current_user_for_VideoCall(UserID),
            relief="flat"
        )
        button_3.place(
            x=2.0,
            y=67.0,
            width=32.0,
            height=32.0
        )

        button_image_4 = PhotoImage(
            file=relative_to_assets("button_16.png"))
        button_4 = Button(
            image=button_image_4,
            borderwidth=0,
            highlightthickness=0,
            command=lambda: get_current_ip_address(UserID),
            relief="flat"
        )
        button_4.place(
            x=2.0,
            y=99.0,
            width=32.0,
            height=32.0
        )

        button_image_5 = PhotoImage(
            file=relative_to_assets("button_17.png"))
        button_5 = Button(
            image=button_image_5,
            borderwidth=0,
            highlightthickness=0,
            command=lambda: TaskAssignmentTool(UserID,Company_Name),
            relief="flat"
        )
        button_5.place(
            x=2.0,
            y=131.0,
            width=32.0,
            height=32.0
        )

        button_image_6 = PhotoImage(
            file=relative_to_assets("button_18.png"))
        button_6 = Button(
            image=button_image_6,
            borderwidth=0,
            highlightthickness=0,
            command=lambda: print("button_18 clicked"),
            relief="flat"
        )
        button_6.place(
            x=2.0,
            y=165.0,
            width=32.0,
            height=32.0
        )

        button_image_7 = PhotoImage(
            file=relative_to_assets("button_19.png"))
        button_7 = Button(
            image=button_image_7,
            borderwidth=0,
            highlightthickness=0,
            command=lambda: print("button_19 clicked"),
            relief="flat"
        )
        button_7.place(
            x=2.0,
            y=197.0,
            width=32.0,
            height=32.0
        )

        button_image_8 = PhotoImage(
            file=relative_to_assets("button_20.png"))
        button_8 = Button(
            image=button_image_8,
            borderwidth=0,
            highlightthickness=0,
            command=lambda: LeaveApplicationDisplay(Company_Name,UserID),
            relief="flat"
        )
        button_8.place(
            x=2.0,
            y=230.0,
            width=32.0,
            height=32.0
        )

        button_image_9 = PhotoImage(
            file=relative_to_assets("button_21.png"))
        button_9 = Button(
            image=button_image_9,
            borderwidth=0,
            highlightthickness=0,
            command=lambda: print("button_21 clicked"),
            relief="flat"
        )
        button_9.place(
            x=2.0,
            y=263.0,
            width=32.0,
            height=32.0
        )

        button_image_10 = PhotoImage(
            file=relative_to_assets("button_22.png"))
        button_10 = Button(
            image=button_image_10,
            borderwidth=0,
            highlightthickness=0,
            command=lambda: AppointmentDisplay(Company_Name,UserID),
            relief="flat"
        )
        button_10.place(
            x=2.0,
            y=298.0,
            width=32.0,
            height=32.0
        )

        button_image_11 = PhotoImage(
            file=relative_to_assets("button_23.png"))
        button_11 = Button(
            image=button_image_11,
            borderwidth=0,
            highlightthickness=0,
            command=lambda:logout(),
            relief="flat"
        )
        button_11.place(
            x=2.0,
            y=412.0,
            width=32.0,
            height=32.0
        )

        canvas.create_rectangle(
            557.0,
            248.0,
            688.0,
            330.0,
            fill="#173054",
            outline="")

        canvas.create_rectangle(
            557.0,
            362.0,
            688.0,
            444.0,
            fill="#173054",
            outline="")

        button_image_12 = PhotoImage(
            file=relative_to_assets("button_24.png"))
        button_12 = Button(
            image=button_image_12,
            borderwidth=0,
            highlightthickness=0,
            command=lambda: print("button_24 clicked"),
            relief="flat"
        )
        button_12.place(
            x=570.0,
            y=187.0,
            width=106.0,
            height=20.0
        )
        
        if image:
            photo = ImageTk.PhotoImage(image)
            canvas.create_image(
                569.0,
                23.0,
                anchor="nw",
                image=photo,
                )
            canvas.image = photo 
        canvas.create_text(
        569.0,
        20.0,
        anchor="nw",
        text="",
        fill="#FFFFFF",
        font=("Inter", 12 * -1)
        )

        canvas.create_text(
            557.0,
            364.0,
            anchor="nw",
            text="task 1\n\n",
            fill="#FFFFFF",
            font=("LibreCaslonText Regular", 12 * -1)
        )

        canvas.create_text(
            557.0,
            379.0,
            anchor="nw",
            text="task 1\n\n",
            fill="#FFFFFF",
            font=("LibreCaslonText Regular", 12 * -1)
        )

        canvas.create_text(
            557.0,
            395.0,
            anchor="nw",
            text="task 1\n\n",
            fill="#FFFFFF",
            font=("LibreCaslonText Regular", 12 * -1)
        )

        canvas.create_text(
            557.0,
            411.0,
            anchor="nw",
            text="task 1\n\n",
            fill="#FFFFFF",
            font=("LibreCaslonText Regular", 12 * -1)
        )

        canvas.create_text(
            557.0,
            426.0,
            anchor="nw",
            text="task 1\n\n",
            fill="#FFFFFF",
            font=("LibreCaslonText Regular", 12 * -1)
        )

def AdminSignInAuthentication(CompanyName,Username,Password):
    Company_Name = CompanyName.get()
    Admin_UserName = Username.get()
    Admin_Password = Password.get()
    print(Company_Name)
    print(Admin_UserName)
    print(Admin_Password)
    
    query = {"CompanyName": Company_Name}
    matching_documents = AdminCollection.find(query)

    for document in matching_documents:
        if document["UserName"] == Admin_UserName and document["Password"] == Admin_Password:
            print("Authentication successful.")
            
            AdminProfile(Admin_UserName,Company_Name)
            # You can perform further actions here if authentication is successful
            break
    else:
        print("Authentication failed. Invalid credentials.")

def WorkspaceSignIn(entry_1):
    global background_image,button_image_1
    canvas = Canvas(
        window,
        bg = "#3A868F",
        height = 450,
        width = 700,
        bd = 0,
        highlightthickness = 0,
        relief = "ridge"
    )

    canvas.place(x = 0, y = 0)
    background_image = PhotoImage(
        file=relative_to_assets("image_14.png"))
    bgimage_1 = canvas.create_image(
        350.0,
        225.0,
        image=background_image
    )

    canvas.create_text(
        415.0,
        72.0,
        anchor="nw",
        text="Welcome Back ",
        fill="#FFFFFF",
        font=("LibreCaslonText Regular", 12 * -1)
    )

    canvas.create_text(
        415.0,
        106.0,
        anchor="nw",
        text="Sign In",
        fill="#FFFFFF",
        font=("LibreCaslonText Bold", 19 * -1)
    )

    canvas.create_text(
        415.0,
        156.0,
        anchor="nw",
        text="Username",
        fill="#FFFFFF",
        font=("LibreCaslonText Regular", 15 * -1)
    )

    canvas.create_text(
        415.0,
        231.0,
        anchor="nw",
        text="Password",
        fill="#FFFFFF",
        font=("LibreCaslonText Regular", 15 * -1)
    )

    entry_image_1 = PhotoImage(
        file=relative_to_assets("entry_31.png"))
    entry_bg_1 = canvas.create_image(
        523.5,
        187.0,
        image=entry_image_1
    )
    UserNameEntry = Entry(
        bd=0,
        bg="#FFFFFF",
        fg="#000716",
        highlightthickness=0
    )
    UserNameEntry.place(
        x=414.0,
        y=173.0,
        width=219.0,
        height=26.0
    )

    entry_image_2 = PhotoImage(
        file=relative_to_assets("entry_32.png"))
    entry_bg_2 = canvas.create_image(
        523.5,
        261.0,
        image=entry_image_2
    )
    PasswordEntry = Entry(
        bd=0,
        bg="#FFFFFF",
        fg="#000716",
        highlightthickness=0
    )
    PasswordEntry.place(
        x=414.0,
        y=247.0,
        width=219.0,
        height=26.0
    )

    button_image_1 = PhotoImage(
        file=relative_to_assets("button_30.png"))
    button_1 = Button(
        image=button_image_1,
        borderwidth=0,
        highlightthickness=0,
        command=lambda: AdminSignInAuthentication(entry_1,UserNameEntry,PasswordEntry),
        relief="flat"
    )
    button_1.place(
        x=468.0,
        y=321.0,
        width=112.0,
        height=36.0
    )

def Workspace():
    global image_image_1,CreateAWorkspaceButton_image_1,entry_image_1,Nextbutton_image_2

    canvas = Canvas(
    window,
    bg = "#ECECD9",
    height = 450,
    width = 700,
    bd = 0,
    highlightthickness = 0,
    relief = "ridge"
    )

    canvas.place(x = 0, y = 0)
    image_image_1 = PhotoImage(
        file=relative_to_assets("image_8.png"))
    image_1 = canvas.create_image(
        350.0,
        231.0,
        image=image_image_1
    )

    canvas.create_text(
        89.0,
        74.0,
        anchor="nw",
        text="Create a new skive\nworkspace",
        fill="#225777",
        font=("Libre Caslon Text", 25 * -1)
    )

    canvas.create_text(
        89.0,
        146.0,
        anchor="nw",
        text="Skive gives your team a home- a place\nwhere they can talk and work together. To\ncreate a new workspace. Click the button\nbelow",
        fill="#225777",
        font=("Libre Caslon Text", 11 * -1)
    )

    CreateAWorkspaceButton_image_1 = PhotoImage(
        file=relative_to_assets("button_11.png"))
    CreateAWorkspaceButton = Button(
        image=CreateAWorkspaceButton_image_1,
        borderwidth=0,
        highlightthickness=0,
        command=Workspacesignup,
        relief="flat"
    )
    CreateAWorkspaceButton.place(
        x=88.0,
        y=225.0,
        width=186.0,
        height=31.0
    )
    
    Nextbutton_image_2 = PhotoImage(
    file=relative_to_assets("button_29.png"))
    Nextbutton_2 = Button(
        image=Nextbutton_image_2,
        borderwidth=0,
        highlightthickness=0,
        command=lambda:WorkspaceSignIn(entry_1),
        relief="flat"
    )
    Nextbutton_2.place(
        x=305.0,
        y=386.0,
        width=90.0,
        height=29.0
    )

    canvas.create_text(
        337.0,
        280.0,
        anchor="nw",
        text="Or",
        fill="#173054",
        font=("Libre Caslon Text", 15 * -1)
    )

    canvas.create_text(
        273.0,
        308.0,
        anchor="nw",
        text="Open a Workspace",
        fill="#173054",
        font=("Libre Caslon Text", 15 * -1)
    )

    entry_image_1 = PhotoImage(
        file=relative_to_assets("entry_7.png"))
    entry_bg_1 = canvas.create_image(
        349.5,
        360.5,
        image=entry_image_1
    )
    
    allrecord= [doc["CompanyName"] for doc in AdminCollection.find({},{'_id':0,'CompanyName':1})]
    CompanyNames=[]
    for prv in allrecord:
        CompanyNames.append(prv)

    print(CompanyNames)
    entry_1 = ttk.Combobox(
        canvas,
        values=CompanyNames
    )
    entry_1.place(
        x=217.0,
        y=346.0,
        width=265.0,
        height=27.0
    )

def receive_messages():
    while True:
        try:
            data = client_socket.recv(1024).decode('utf-8')
            message_listbox.insert(tk.END, data)
        except ConnectionAbortedError:
            break

def send_message(my_message,selected_user,recipients_ip_address):
    message = my_message.get()
    if message:
        recipient_ip = recipients_ip_address
        full_message = f"{recipient_ip}:{message}"
        client_socket.send(full_message.encode('utf-8'))
        message_listbox.insert(tk.END, "You: " + message)
        my_message.set("")

def connect_to_server(selected_user,recipient_ipaddress):
    global client_socket
    global receive_thread
    employee = EmployeeCollection.find_one({'UserName': selected_user})
    if employee:
            server_ip = employee.get('server_ip', "IP address not found")
            server_port = int(employee.get('server_port', "IP address not found"))
            print(server_ip)
            print(server_port)
            PersonalChatGUI(selected_user,recipient_ipaddress)
            
        
    else:
        return "server not found"
    
    client_socket = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
    client_socket.connect((server_ip, server_port))
    receive_thread = threading.Thread(target=receive_messages)
    receive_thread.start()
  
def button_click(selected_user):
        employee = EmployeeCollection.find_one({'UserName': selected_user})
        DirectMessageframe.destroy()
        DirectMessageframe.pack_forget()
        window.update()
        if employee:
            ip_address = employee.get('ip_address', "IP address not found")
            print("this will display the selected user and their ip address")
            print(selected_user)
            print(ip_address)
            connect_to_server(selected_user,ip_address)
            
        
        else:
            return "Employee not found"
        
        
        #print("Selected user:", ip_address)

def get_current_ip_address(username):
    ip_address = socket.gethostbyname(socket.gethostname())
    
    # Check if the username already exists in the collection
    existing_user = EmployeeCollection.find_one({"UserName": username})
    if existing_user:
        # Update the existing document with the new IP address
        EmployeeCollection.update_one({"UserName": username}, {"$set": {"ip_address": ip_address}})
        Display_current_user_for_direct_message(username)

def get_ip_address_of_invited_users(invited_usernames):
    ip_addresses = []
    for username in invited_usernames:
        employee = EmployeeCollection.find_one({'UserName': username})
        if employee:
            ip_address = employee.get('ip_address', "IP address not found")
            ip_addresses.append(ip_address)
        else:
            print(f"IP address for {username} not found in Employee collection.")
    return ip_addresses

def display_notification():
    # Function to display the notification window
    window = tk.Tk()
    window.title("Notification")
    
    lbl_notification = tk.Label(window, text="Join Meeting")
    lbl_notification.pack()
    
    btn_join = tk.Button(window, text="Join", command=join_meeting)
    btn_join.pack()
    
    window.mainloop()

def JoinMeetingHost(username, invited_user_ipaddress):
   
    host_meeting_window = tk.Tk()
    host_meeting_window.title('Host Meeting')
    host_meeting_window.geometry('800x600')
    local_ip_address = socket.gethostbyname(socket.gethostname())

    server = StreamingServer(local_ip_address, 9999)
    receiver = AudioReceiver(local_ip_address, 8888)

    camera_clients = []
    screen_share_clients = []
    audio_senders = []

    def start_listening():
        t1 = threading.Thread(target=server.start_server)
        t2 = threading.Thread(target=receiver.start_server)
        t1.start()
        t2.start()

    def start_camera_stream(target_ips):
        for ip in target_ips:
            camera_client = CameraClient(ip.strip(), 7777)
            camera_clients.append(camera_client)
            t3 = threading.Thread(target=camera_client.start_stream)
            t3.start()

    def start_screen_sharing(target_ips):
        for ip in target_ips:
            screen_client = ScreenShareClient(ip.strip(), 7777)
            screen_share_clients.append(screen_client)
            t4 = threading.Thread(target=screen_client.start_stream)
            t4.start()

    def start_audio_stream(target_ips):
        for ip in target_ips:
            audio_sender = AudioSender(ip.strip(), 6666)
            audio_senders.append(audio_sender)
            t5 = threading.Thread(target=audio_sender.start_stream)
            t5.start()

    def stop_camera_stream():
        for client in camera_clients:
            client.stop_stream()

    def stop_screen_sharing():
        for client in screen_share_clients:
            client.stop_stream()

    def stop_audio_stream():
        for sender in audio_senders:
            sender.stop_stream()

    # def start_conferencing(target_ips):
    #     start_camera_stream(target_ips)
    #     start_screen_sharing(target_ips)
    #     start_audio_stream(target_ips)

    def toggle_camera_stream():
        if btn_camera.config('text')[-1] == 'Start Camera Stream':
            btn_camera.config(text='Stop Camera Stream')
            start_camera_stream(invited_user_ipaddress)
        else:
            btn_camera.config(text='Start Camera Stream')
            stop_camera_stream()

    def toggle_screen_sharing():
        if btn_screen.config('text')[-1] == 'Start Screen Sharing':
            btn_screen.config(text='Stop Screen Sharing')
            start_screen_sharing(invited_user_ipaddress)
        else:
            btn_screen.config(text='Start Screen Sharing')
            stop_screen_sharing()

    def toggle_audio_stream():
        if btn_audio.config('text')[-1] == 'Start Audio Stream':
            btn_audio.config(text='Stop Audio Stream')
            start_audio_stream(invited_user_ipaddress)
        else:
            btn_audio.config(text='Start Audio Stream')
            stop_audio_stream()

    # GUI
    window = tk.Tk()
    window.title('Video Conference')
    window.geometry('800x600')

    label_target_ip = tk.Label(window, text="Target IPs (separated by comma):")
    label_target_ip.pack()

    btn_listen = tk.Button(window, text="Start Listening", width=20, command=start_listening(invited_user_ipaddress))
    btn_listen.pack(anchor=tk.CENTER, expand=True)

    btn_camera = tk.Button(window, text="Start Camera Stream", width=20, command=toggle_camera_stream(invited_user_ipaddress))
    btn_camera.pack(side=tk.LEFT, padx=10)

    btn_screen = tk.Button(window, text="Start Screen Sharing", width=20, command=toggle_screen_sharing(invited_user_ipaddress))
    btn_screen.pack(side=tk.LEFT, padx=10)

    btn_audio = tk.Button(window, text="Start Audio Stream", width=20, command=toggle_audio_stream(invited_user_ipaddress))
    btn_audio.pack(side=tk.LEFT, padx=10)

    # btn_start_conference = tk.Button(window, text="Start Video Conference", width=20, command=start_conferencing(invited_user_ipaddress))
    # btn_start_conference.pack(anchor=tk.CENTER, expand=True)

    window.mainloop()

def JoinVideoConferencing(invited_user_ipaddress,UserName):
    global TurnOnMicButton, TurnOnCameraButton, JoinMeetingButton, JoinMeetingBackgroundImage, camera_on, mic_on,TurnOnMicOffImage,TurnOnMicOnImage,TurnOnCameraOffImage,TurnOnCameraOnImage

    def send_notification(invited_user_ipaddress):
        document = EmployeeCollection.find_one({})
        SERVER_IP = document.get("NotificationServerIP")
        SERVER_PORT = document.get("NotificationServerPort")
        static_message = f"Join the Meeting. You are invited by {UserName}"
        
        if not invited_user_ipaddress:
            print("No recipient IP addresses provided")
            return

        try:
            for recipient_ip in invited_user_ipaddress:
                print(recipient_ip)
                client_socket = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
                client_socket.connect((SERVER_IP, SERVER_PORT))
                client_socket.send(f"{recipient_ip}:{static_message}".encode())
                client_socket.close()
            print("Notification sent successfully")
        except Exception as e:
            print(f"Failed to send notification: {e}")    
    
    send_notification(invited_user_ipaddress)
    def toggle_camera():
        global camera_on
        if camera_on:
            # Turn off camera
            
            # You can add the code to turn off the camera here
            print("Camera turned off")
            TurnOnCameraButton.config(image=TurnOnCameraOnImage)  # Change button image
            camera_on = False
        else:
            # Turn on camera
            
            # You can add the code to turn on the camera here
            print("Camera turned on")
            TurnOnCameraButton.config(image=TurnOnCameraOffImage)  # Change button image
            camera_on = True

    def toggle_microphone():
        global mic_on
        if mic_on:
            # Turn off microphone
            # You can add the code to turn off the microphone here
            print("Microphone turned off")
            TurnOnMicButton.config(image=TurnOnMicOnImage)  # Change button image
            mic_on = False
        else:
            # Turn on microphone
            # You can add the code to turn on the microphone here
            print("Microphone turned on")
            TurnOnMicButton.config(image=TurnOnMicOffImage)  # Change button image
            mic_on = True

    # Destroy previous frame
    DisplayCurrentUserFrameVideoCall.destroy()
    

    # Create new frame for joining video conference
    JoinVideoConferenceFrame = Frame(window, width=510, height=450)
    JoinVideoConferenceFrame.place(x=36, y=0)
    canvas = Canvas(
        JoinVideoConferenceFrame,
        bg="#3A868F",
        height=450,
        width=510,
        bd=0,
        highlightthickness=0,
        relief="ridge"
    )
    canvas.place(x=0, y=0)

    JoinMeetingBackgroundImage = PhotoImage(file=relative_to_assets("image_20.png"))
    image_1 = canvas.create_image(
        255.0,
        225.0,
        image=JoinMeetingBackgroundImage
    )

    canvas.create_rectangle(
        42.0,
        20.0,
        468.0,
        293.0,
        fill="#ECECD9",
        outline="")

    # Load images for camera and microphone buttons
    TurnOnCameraOnImage = PhotoImage(file=relative_to_assets("button_40.png"))
    TurnOnCameraOffImage = PhotoImage(file=relative_to_assets("button_43.png"))
    TurnOnMicOnImage = PhotoImage(file=relative_to_assets("button_41.png"))
    TurnOnMicOffImage = PhotoImage(file=relative_to_assets("button_44.png"))

    # Set initial state for camera and microphone
    camera_on = False
    mic_on = False

    # Create camera button
    TurnOnCameraButton = Button(
        canvas,
        image=TurnOnCameraOnImage,  # Initial state is off
        borderwidth=0,
        highlightthickness=0,
        command=toggle_camera,  # Bind toggle_camera function to button click event
        relief="flat"
    )
    TurnOnCameraButton.place(
        x=216.0,
        y=307.0,
        width=37.0,
        height=37.0
    )

    # Create microphone button
    TurnOnMicButton = Button(
        canvas,
        image=TurnOnMicOnImage,  # Initial state is off
        borderwidth=0,
        highlightthickness=0,
        command=toggle_microphone,  # Bind toggle_microphone function to button click event
        relief="flat"
    )
    TurnOnMicButton.place(
        x=257.0,
        y=307.0,
        width=37.0,
        height=37.0
    )

    JoinMeetingButton = PhotoImage(file=relative_to_assets("button_42.png"))
    button_3 = Button(
        canvas,
        image=JoinMeetingButton,
        borderwidth=0,
        highlightthickness=0,
        command=lambda:JoinMeetingHost(UserName,invited_user_ipaddress),
        relief="flat"
        )
    button_3.place(
        x=210.0,
        y=374.0,
        width=94.0,
        height=35.0
    )

    print("----------------------")
    print(invited_users_ipaddress)

def InvitedUsersForVideoConferencing(username):
    global invited_users,invited_users_ipaddress
    
    # Check if the user is already invited
    if username not in invited_users:
        NextButtonFrame = Frame(window,height=450, width=510)
        NextButtonFrame.place(x=36 ,y=400)
        employee = EmployeeCollection.find_one({'UserName': username})
        invited_users.append(username)
        ip_address = employee.get('ip_address', "IP address not found")
        print(f"User: {username}, IP Address: {ip_address}")
        invited_users_ipaddress.append(ip_address)
        print(f"{username} invited.")
        print(invited_users_ipaddress)
    else:
        print(f"{username} is already invited.")

def Display_current_user_for_VideoCall(username):
    global VideoCallframe, VideoCallNextButton,DisplayCurrentUserFrameVideoCall
    DisplayCurrentUserFrameVideoCall = Frame(window,height=450, width=510)
    DisplayCurrentUserFrameVideoCall.place(x=36, y=0)
    canvas = Canvas(
    DisplayCurrentUserFrameVideoCall,
    bg = "#3A868F",
    height = 450,
    width = 510,
    bd = 0,
    highlightthickness = 0,
    relief = "ridge"
    )

    canvas.place(x = 0, y = 0)
    
    VideoCallNextButton = PhotoImage(
        file=relative_to_assets("button_39.png"))
    button_1 = Button(
        image=VideoCallNextButton,
        borderwidth=0,
        highlightthickness=0,
        command=lambda: JoinVideoConferencing(invited_users_ipaddress,username),
        relief="flat"
    )
    button_1.place(
        x=210.0,
        y=374.0,
        width=94.0,
        height=35.0
    )
    VideoCallframe = customtkinter.CTkScrollableFrame(canvas, width=318)
    VideoCallframe.pack(padx=(0, 0))
    VideoCallframe.config(bg="#ECECD9", bd=0)
    VideoCallframe.config(height=500)
    
    # Function to handle button click
    users = [doc["UserName"] for doc in EmployeeCollection.find({"UserName":{"$ne": username}},{'_id':0,'UserName':1})]
    print(users)
    for user in users:
        user_data = EmployeeCollection.find_one({"UserName": user})
        if user_data and "photo" in user_data:
            photo_data_base64 = user_data["photo"]
            # Decode Base64 data
            photo_data_binary = base64.b64decode(photo_data_base64)
            with open("user_photo.jpg", "wb") as f:
                f.write(photo_data_binary)
            image = Image.open("user_photo.jpg")
            image = image.resize((35, 35), Image.LANCZOS)
            image = round_corners(image, 50)
        
        else:
            print("Photo not found for the specified user.")
            
        
        

        user_frame = ttk.Frame(VideoCallframe, padding=5, relief="raised")
        user_frame.pack(fill=tk.BOTH, expand=True)
        
        photo = ImageTk.PhotoImage(image)
        image_label = tk.Label(user_frame, image=photo)
        image_label.image = photo  # Keep a reference to avoid garbage collection
        image_label.pack(anchor="w")
        
        label = ttk.Label(user_frame, text=user)
        label.pack(side=tk.LEFT)
        
        button = ttk.Button(user_frame, text="Invite", command=lambda u=user: InvitedUsersForVideoConferencing(u))
        button.pack(side=tk.RIGHT)

def Display_current_user_for_direct_message(username):
    global DirectMessageframe, VideoCallframe
    
    window.update()
    DirectMessageframe = customtkinter.CTkScrollableFrame(window, width=510, height=450)
    DirectMessageframe.pack(padx=(24, 104))
    DirectMessageframe.config(bg="#ECECD9", bd=0)
    #DirectMessageframe = customtkinter.CTkScrollableFrame(window, width=510, height=450)
    #DirectMessageframe.pack(padx=(36, 154), pady=0)

    #main_frame = ttk.Frame(window, padding=10)
    #main_frame.pack(fill=tk.BOTH, expand=True)
    
    # Function to handle button click
    users = [doc["UserName"] for doc in EmployeeCollection.find({"UserName":{"$ne": username}},{'_id':0,'UserName':1})]
    
    for user in users:
        user_data = EmployeeCollection.find_one({"UserName": user})
        if user_data and "photo" in user_data:
            photo_data_base64 = user_data["photo"]
            # Decode Base64 data
            photo_data_binary = base64.b64decode(photo_data_base64)
            with open("user_photo.jpg", "wb") as f:
                f.write(photo_data_binary)
            image = Image.open("user_photo.jpg")
            image = image.resize((35, 35), Image.LANCZOS)
            image = round_corners(image, 50)
        
        else:
            print("Photo not found for the specified user.")
        user_frame = ttk.Frame(DirectMessageframe, padding=5, relief="raised")
        user_frame.pack(fill=tk.BOTH, expand=True)
        
        photo = ImageTk.PhotoImage(image)
        image_label = tk.Label(user_frame, image=photo)
        image_label.image = photo  # Keep a reference to avoid garbage collection
        image_label.pack(anchor="w")

        label = ttk.Label(user_frame, text=user)
        label.pack(side=tk.LEFT)
        
        button = ttk.Button(user_frame, text="Message", command=lambda u=user: button_click(u))
        button.pack(side=tk.RIGHT)

def MailApplication():
    pass

def PersonalChatGUI(selected_user,recipients_ip_address): 
    DirectMessageframe.pack_forget()
    
    window.update()
    global BackButton,MyMessage,SendMessageButton,message_listbox
    PersonalChatCanvas = Canvas(
    window,
    bg = "#173054",
    height = 450,
    width = 510,
    bd = 0,
    highlightthickness = 0,
    relief = "ridge"
    )

    PersonalChatCanvas.place(x = 36, y = 0)
    PersonalChatCanvas.config(bg="#173054")
    PersonalChatCanvas.create_rectangle(
        0.0,
        0.0,
        510.0,
        40.0,
        fill="#225777",
        outline="")

    BackButton = PhotoImage(
        file=relative_to_assets("button_27.png"))
    button_1 = Button(
        image=BackButton,
        borderwidth=0,
        highlightthickness=0,
        command=lambda: print("Back Button clicked"),
        relief="flat"
    )
    button_1.place(
        x=44.0,
        y=4.0,
        width=32.0,
        height=32.0
    )

    PersonalChatCanvas.create_text(
        52.0,
        7.0,
        anchor="nw",
        text="" + selected_user,
        fill="#ECECD9",
        font=("LibreCaslonText Regular", 20 * -1)
    )

    PersonalChatCanvas.create_rectangle(
        0.0,
        405.0,
        510.0,
        450.0,
        fill="#ECECD9",
        outline="")
    
    message_listbox = Listbox(PersonalChatCanvas, width=50, height=13)
    message_listbox.place(x=1,y=41)

    MyMessage = PhotoImage(
        file=relative_to_assets("entry_24.png"))
    entry_bg_1 = PersonalChatCanvas.create_image(
        199.0,
        428.0,
        image=MyMessage
    )
    my_message = StringVar()
    entry_1 = Entry(
        bd=0,
        bg="#225777",
        fg="#000716",
        highlightthickness=0,
        textvariable=my_message,
    )
    entry_1.place(
        x=50.0,
        y=411.0,
        width=370.0,
        height=32.0
    )
    entry_1.bind("<Return>", send_message)
    entry_1.config(bg="#225777")
    entry_1.config(bd=0)
    entry_1.config(highlightthickness=0)

    SendMessageButton = PhotoImage(
        file=relative_to_assets("button_28.png"))
    button_2 = Button(
        image=SendMessageButton,
        borderwidth=0,
        highlightthickness=0,
        command=lambda: send_message(my_message,selected_user,recipients_ip_address),
        relief="flat"
    )
    button_2.place(
        x=436.0,
        y=411.0,
        width=102.0,
        height=34.0
    )

def check_CompanyName(CompanyName):
    existing_Company = EmployeeCollection.find_one({"CompanyName": CompanyName})
    return existing_Company is not None

def ManagerFormDisplay(CompanyName,UserName,Email,Password):
    global button_image_1,button_image_2,entry_image_1,entry_image_2,entry_image_3,entry_image_4,entry_image_5,entry_image_6,entry_image_7,entry_image_8,entry_image_9
    canvas = Canvas(
    window,
    bg = "#225777",
    height = 450,
    width = 700,
    bd = 0,
    highlightthickness = 0,
    relief = "ridge"
    )

    canvas.place(x = 0, y = 0)
    canvas.create_rectangle(
        57.0,
        41.0,
        642.0,
        408.0,
        fill="#D9D9D9",
        outline="")

    canvas.create_text(
        72.0,
        56.0,
        anchor="nw",
        text="Personal Details",
        fill="#173054",
        font=("LibreCaslonText Regular", 14 * -1)
    )

    canvas.create_text(
        88.0,
        86.0,
        anchor="nw",
        text="First Name",
        fill="#173054",
        font=("LibreCaslonText Regular", 11 * -1)
    )

    canvas.create_text(
        89.0,
        137.0,
        anchor="nw",
        text="Age",
        fill="#173054",
        font=("LibreCaslonText Regular", 11 * -1)
    )

    canvas.create_text(
        89.0,
        187.0,
        anchor="nw",
        text="Birthdate(DD/MM/YYYY)",
        fill="#173054",
        font=("LibreCaslonText Regular", 11 * -1)
    )

    canvas.create_text(
        281.0,
        187.0,
        anchor="nw",
        text="Designation",
        fill="#173054",
        font=("LibreCaslonText Regular", 11 * -1)
    )

    canvas.create_text(
        281.0,
        137.0,
        anchor="nw",
        text="Gender",
        fill="#173054",
        font=("LibreCaslonText Regular", 11 * -1)
    )

    canvas.create_text(
        473.0,
        137.0,
        anchor="nw",
        text="Mobile Number",
        fill="#173054",
        font=("LibreCaslonText Regular", 11 * -1)
    )

    canvas.create_text(
        473.0,
        187.0,
        anchor="nw",
        text="Highest Education",
        fill="#173054",
        font=("LibreCaslonText Regular", 11 * -1)
    )

    canvas.create_text(
        281.0,
        87.0,
        anchor="nw",
        text="Middle Name",
        fill="#173054",
        font=("LibreCaslonText Regular", 11 * -1)
    )

    canvas.create_text(
        474.0,
        86.0,
        anchor="nw",
        text="Last Name",
        fill="#173054",
        font=("LibreCaslonText Regular", 11 * -1)
    )

    entry_image_1 = PhotoImage(
        file=relative_to_assets("entry_33.png"))
    entry_bg_1 = canvas.create_image(
        158.0,
        110.5,
        image=entry_image_1
    )
    ManagerFirstNameEntry = Entry(
        bd=0,
        bg="#173054",
        fg="#000716",
        highlightthickness=0
    )
    ManagerFirstNameEntry.place(
        x=90.0,
        y=101.0,
        width=136.0,
        height=17.0
    )

    entry_image_2 = PhotoImage(
        file=relative_to_assets("entry_34.png"))
    entry_bg_2 = canvas.create_image(
        350.0,
        110.5,
        image=entry_image_2
    )
    ManagerMiddleNameEntry = Entry(
        bd=0,
        bg="#173054",
        fg="#000716",
        highlightthickness=0
    )
    ManagerMiddleNameEntry.place(
        x=282.0,
        y=101.0,
        width=136.0,
        height=17.0
    )

    entry_image_3 = PhotoImage(
        file=relative_to_assets("entry_35.png"))
    entry_bg_3 = canvas.create_image(
        542.0,
        110.5,
        image=entry_image_3
    )
    ManagerLastNameEntry = Entry(
        bd=0,
        bg="#173054",
        fg="#000716",
        highlightthickness=0
    )
    ManagerLastNameEntry.place(
        x=474.0,
        y=101.0,
        width=136.0,
        height=17.0
    )

    entry_image_4 = PhotoImage(
        file=relative_to_assets("entry_36.png"))
    entry_bg_4 = canvas.create_image(
        158.0,
        160.5,
        image=entry_image_4
    )
    ManagerAgeEntry = Entry(
        bd=0,
        bg="#173054",
        fg="#000716",
        highlightthickness=0
    )
    ManagerAgeEntry.place(
        x=90.0,
        y=151.0,
        width=136.0,
        height=17.0
    )

    entry_image_5 = PhotoImage(
        file=relative_to_assets("entry_37.png"))
    entry_bg_5 = canvas.create_image(
        350.0,
        160.5,
        image=entry_image_5
    )
    ManagerGenderEntry = Entry(
        bd=0,
        bg="#173054",
        fg="#000716",
        highlightthickness=0
    )
    ManagerGenderEntry.place(
        x=282.0,
        y=151.0,
        width=136.0,
        height=17.0
    )

    entry_image_6 = PhotoImage(
        file=relative_to_assets("entry_38.png"))
    entry_bg_6 = canvas.create_image(
        542.0,
        160.5,
        image=entry_image_6
    )
    ManagerMobileNumberEntry = Entry(
        bd=0,
        bg="#173054",
        fg="#000716",
        highlightthickness=0
    )
    ManagerMobileNumberEntry.place(
        x=474.0,
        y=151.0,
        width=136.0,
        height=17.0
    )

    entry_image_7 = PhotoImage(
        file=relative_to_assets("entry_39.png"))
    entry_bg_7 = canvas.create_image(
        158.0,
        210.5,
        image=entry_image_7
    )
    ManagerBirthDateEntry = Entry(
        bd=0,
        bg="#173054",
        fg="#000716",
        highlightthickness=0
    )
    ManagerBirthDateEntry.place(
        x=90.0,
        y=201.0,
        width=136.0,
        height=17.0
    )

    entry_image_8 = PhotoImage(
        file=relative_to_assets("entry_40.png"))
    entry_bg_8 = canvas.create_image(
        350.0,
        210.5,
        image=entry_image_8
    )
    ManagerDesignationEntry = Text(
        bd=0,
        bg="#173054",
        fg="#000716",
        highlightthickness=0
    )
    ManagerDesignationEntry.place(
        x=282.0,
        y=201.0,
        width=136.0,
        height=17.0
    )

    entry_image_9 = PhotoImage(
        file=relative_to_assets("entry_41.png"))
    entry_bg_9 = canvas.create_image(
        542.0,
        210.5,
        image=entry_image_9
    )
    ManagerEducationEntry = Entry(
        bd=0,
        bg="#173054",
        fg="#000716",
        highlightthickness=0
    )
    ManagerEducationEntry.place(
        x=474.0,
        y=201.0,
        width=136.0,
        height=17.0
    )

    button_image_1 = PhotoImage(
        file=relative_to_assets("button_31.png"))
    button_1 = Button(
        image=button_image_1,
        borderwidth=0,
        highlightthickness=0,
        command=lambda: Managerbrowse_photo(),
        relief="flat"
    )
    button_1.place(
        x=113.0,
        y=330.0,
        width=118.0,
        height=28.0
    )

    button_image_2 = PhotoImage(
        file=relative_to_assets("button_32.png"))
    button_2 = Button(
        image=button_image_2,
        borderwidth=0,
        highlightthickness=0,
        command=lambda: ManagerSubmitForm(CompanyName,UserName,Email,Password,ManagerFirstNameEntry,ManagerMiddleNameEntry,ManagerLastNameEntry,ManagerAgeEntry,ManagerGenderEntry,ManagerMobileNumberEntry,ManagerBirthDateEntry,ManagerDesignationEntry,ManagerEducationEntry),
        relief="flat"
    )
    button_2.place(
        x=469.0,
        y=330.0,
        width=118.0,
        height=28.0
    )
    
def ManagerSubmitForm(CompanyName,UserName,Email,Password,ManagerFirstNameEntry,ManagerMiddleNameEntry,ManagerLastNameEntry,ManagerAgeEntry,ManagerGenderEntry,ManagerMobileNumberEntry,ManagerBirthDateEntry,ManagerDesignationEntry,ManagerEducationEntry):
    photo_path = Managerbrowse_photo()
    First_Name = ManagerFirstNameEntry.get()
    Middle_Name = ManagerMiddleNameEntry.get()
    Last_Name = ManagerLastNameEntry.get()
    Age = ManagerAgeEntry.get()
    Birthdate = ManagerBirthDateEntry.get()
    Designation = ManagerDesignationEntry.get("1.0", "end-1c")
    Gender = ManagerGenderEntry.get()
    Mobile_Number = ManagerMobileNumberEntry.get()
    Highest_Education = ManagerEducationEntry.get()
    
    photo_path = photo_path  # Get the path of the uploaded photo
    if CompanyName and UserName and Password and First_Name and Middle_Name and Last_Name and Age and Birthdate and Designation and Gender and Mobile_Number and Highest_Education and photo_path:
        with open(photo_path, "rb") as photo_file:
            photo_data = base64.b64encode(photo_file.read()).decode('utf-8')
            data = {
                "CompanyName" : CompanyName,
                "UserName" : UserName,
                "Password" : Password,
                "First_Name" : First_Name,
                "Middle_Name" : Middle_Name,
                "Last_Name" : Last_Name,
                "Age": Age,
                "Birthdate" : Birthdate,
                "Designation" : Designation,
                "Gender" : Gender,
                "Mobile_Number" : Mobile_Number,
                "Highest_Education" : Highest_Education,
                "photo": photo_data  # Store the photo data directly in the document
            }
            AdminCollection.insert_one(data)
            print("Data saved successfully!")
            
            AdminProfile(UserName,CompanyName)

def WorkspaceSignUpAuthentication(entry_2,entry_3,entry_4,entry_5, entry_6):
    
    CompanyName = entry_2.get()
    UserName = entry_3.get()
    Email = entry_4.get()
    Password = entry_5.get()
    ConfirmPassword = entry_6.get()
    #CompanyDBName = CompanyName.replace(" ","")

    if check_CompanyName(CompanyName):
        messagebox.showerror("Error", "Company already exists!")
        return
    
    if Password == ConfirmPassword:
        messagebox.showinfo('Successful Sign Up ','Done') 
        ManagerFormDisplay(CompanyName,UserName,Email,Password)
        #CompanyDatabase = client["" + CompanyDBName]
        #CollectionName = f"{CompanyDBName}_Collection"
        #AdminCollection = CompanyDatabase[CollectionName]
        '''Workspace_manager = {
        "CompanyName": CompanyName,
        "UserName" : UserName,
        "Email": Email,
        "Password": Password
        }  
        res=AdminCollection.insert_one(Workspace_manager)
        
        if res.acknowledged:
            print(f"Data inserted with ObjectID: {res.inserted_id}")
            ManagerFormDisplay(CompanyName)
        else:
            print("Failed to insert data")  
        
    else:
        messagebox.showerror('Invalid',"Passwords don't match")'''
    
def logout():
    Main_Window()

def Workspacesignup():
    global image_image_1,entry_image_1,entry_image_2,entry_image_3,entry_image_4,entry_image_5,entry_image_6,button_image_1
    canvas = Canvas(
    window,
    bg = "#173054",
    height = 450,
    width = 700,
    bd = 0,
    highlightthickness = 0,
    relief = "ridge"
    )

    canvas.place(x = 0, y = 0)
    image_image_1 = PhotoImage(
        file=relative_to_assets("image_12.png"))
    image_1 = canvas.create_image(
        178.0,
        225.0,
        image=image_image_1
    )
    canvas.config(bg="#173054")

    button_image_1 = PhotoImage(
        file=relative_to_assets("button_12.png"))
    CreateAWorkspaceSubmitbutton_1 = Button(
        image=button_image_1,
        borderwidth=0,
        highlightthickness=0,
        command=lambda: WorkspaceSignUpAuthentication(entry_2, entry_3, entry_4, entry_5, entry_6),
        relief="flat"
    )
    CreateAWorkspaceSubmitbutton_1.place(
        x=436.0,
        y=392.0,
        width=147.0,
        height=33.0
    )

    canvas.create_text(
        381.0,
        76.0,
        anchor="nw",
        text="Enter Company Name",
        fill="#ECECD9",
        font=("Libre Caslon Text", 14 * -1)
    )

    canvas.create_text(
        381.0,
        26.0,
        anchor="nw",
        text="Welcome to \nSkive",
        fill="#ECECD9",
        font=("LibreCaslonText Bold", 20 * -1)
    )

    canvas.create_text(
        381.0,
        131.0,
        anchor="nw",
        text="Create Username",
        fill="#ECECD9",
        font=("Libre Caslon Text", 14 * -1)
    )

    canvas.create_text(
        381.0,
        181.0,
        anchor="nw",
        text="Enter your email",
        fill="#ECECD9",
        font=("Libre Caslon Text", 14 * -1)
    )

    canvas.create_text(
        381.0,
        233.0,
        anchor="nw",
        text="Password",
        fill="#ECECD9",
        font=("Libre Caslon Text", 14 * -1)
    )

    canvas.create_text(
        381.0,
        286.0,
        anchor="nw",
        text="Confirm password",
        fill="#ECECD9",
        font=("Libre Caslon Text", 14 * -1)
    )

    
    entry_image_2 = PhotoImage(
        file=relative_to_assets("entry_26.png"))
    entry_bg_2 = canvas.create_image(
        510.5,
        106.0,
        image=entry_image_2
    )
    entry_2 = Entry(
        bd=0,
        bg="#ECECD9",
        fg="#000716",
        highlightthickness=0
    )
    entry_2.place(
        x=384.0,
        y=94.0,
        width=253.0,
        height=22.0
    )

    entry_image_3 = PhotoImage(
        file=relative_to_assets("entry_27.png"))
    entry_bg_3 = canvas.create_image(
        509.5,
        160.0,
        image=entry_image_3
    )
    entry_3 = Entry(
        bd=0,
        bg="#ECECD9",
        fg="#000716",
        highlightthickness=0
    )
    entry_3.place(
        x=383.0,
        y=148.0,
        width=253.0,
        height=22.0
    )

    entry_image_4 = PhotoImage(
        file=relative_to_assets("entry_28.png"))
    entry_bg_4 = canvas.create_image(
        510.5,
        211.0,
        image=entry_image_4
    )
    entry_4 = Entry(
        bd=0,
        bg="#ECECD9",
        fg="#000716",
        highlightthickness=0
    )
    entry_4.place(
        x=384.0,
        y=199.0,
        width=253.0,
        height=22.0
    )

    entry_image_5 = PhotoImage(
        file=relative_to_assets("entry_29.png"))
    entry_bg_5 = canvas.create_image(
        510.5,
        262.0,
        image=entry_image_5
    )
    entry_5 = Entry(
        bd=0,
        bg="#ECECD9",
        fg="#000716",
        highlightthickness=0
    )
    entry_5.place(
        x=384.0,
        y=250.0,
        width=253.0,
        height=22.0
    )

    entry_image_6 = PhotoImage(
        file=relative_to_assets("entry_30.png"))
    entry_bg_6 = canvas.create_image(
        510.5,
        315.0,
        image=entry_image_6
    )
    entry_6 = Entry(
        bd=0,
        bg="#ECECD9",
        fg="#000716",
        highlightthickness=0
    )
    entry_6.place(
        x=384.0,
        y=303.0,
        width=253.0,
        height=22.0
    )
 
def round_corners(image, radius):
    # Create a mask for the rounded corners
    mask = Image.new("L", image.size, 0)
    draw = ImageDraw.Draw(mask)
    draw.rounded_rectangle((0, 0, image.width, image.height), radius, fill=255)
    
    # Apply the mask to the image
    result = Image.new("RGBA", image.size)
    result.paste(image, mask=mask)
    return result

def GroupChatApplication(CompanyName,UserID):
    global LeftFramePhotoImageGroupChat,CreateGroupChatButtonImage,GroupChatFrame
    # Main frame
    GroupChatFrame = tk.Frame(window, width=510, height=450)
    GroupChatFrame.place(x=36,y=0)
    # Left frame for group chat label and search button
    # left_frame = tk.Frame(GroupChatFrame, width=190, height=450)
    # left_frame.place(x=0,y=0)

    created_groups = []
    def DisplayGroups():
        # Assuming GroupChatCollection is a collection object connected to MongoDB

        # Query documents from GroupChatCollection
        group_chats = GroupChatCollection.find({
        "$or": [
            {"CreatedBy": UserID},
            {"SelectedUsers": UserID}
        ]
    })

        # Variables for positioning widgets
        y_offset = 45  # Initial y-offset for placing widgets
        distance_between_rectangles = 10  # Adjust this value to set the distance between rectangles

        # Function to execute when a rectangle is clicked
        def on_rectangle_click(event, group_name):
            global RightFrameBGImage,GroupChatMessageEntry,GroupChatSendMessageButton
            # You can use the group name passed as an argument
            def on_send_message():
                # Retrieve the message from the entry widget
                message = entry_1.get()
                sender = UserID  # Assuming the sender's username is fixed for now
                message_listbox.insert(tk.END, f"{sender}: {message}")
                
                

                # Store the message in the database
                if GroupChatCollection.find_one({"GroupName": group_name}):
                    # Group name exists, update the existing document
                    GroupChatCollection.update_one(
                        {"GroupName": group_name},
                        {"$push": {
                            "Messages": {
                                "Sender": sender,
                                "Content": message,
                                "Timestamp": datetime.now()
                            }
                        }}
                    )
                    

                

                # Check if the group exists in the database
                
            canvas = Canvas(
                right_frame,
                bg = "#173054",
                height = 450,
                width = 320,
                bd = 0,
                highlightthickness = 0,
                relief = "ridge"
            )

            canvas.place(x = 0, y = 0)
            RightFrameBGImage = PhotoImage(
                file=relative_to_assets("image_19.png"))
            image_1 = canvas.create_image(
                160.0,
                225.0,
                image=RightFrameBGImage
            )

            canvas.create_rectangle(
                0.0,
                0.0,
                320.0,
                40.0,
                fill="#ECECD9",
                outline="")

            canvas.create_rectangle(
                0.0,
                407.0,
                320.0,
                450.0,
                fill="#ECECD9",
                outline="")
            

            message_listbox = Listbox(canvas, width=50, height=14)
            message_listbox.place(x=0,y=40)
          
            
            def updateListbox(group_name):
                group_messages = GroupChatCollection.find_one({"GroupName": group_name})
                if group_messages and "Messages" in group_messages:
                    for message in group_messages["Messages"]:
                        sender = message["Sender"]
                        content = message["Content"]
                        message_listbox.insert(tk.END, f"{sender}: {content}")

            GroupChatMessageEntry = PhotoImage(
                file=relative_to_assets("entry_55.png"))
            entry_bg_1 = canvas.create_image(
                140.5,
                428.5,
                image=GroupChatMessageEntry
            )
            entry_1 = Entry(
                canvas,
                bd=0,
                bg="#225777",
                fg="#000716",
                highlightthickness=0
            )
            entry_1.place(
                x=16.0,
                y=414.0,
                width=249.0,
                height=27.0
            )
            updateListbox(group_name)

            GroupChatSendMessageButton = PhotoImage(
                file=relative_to_assets("button_38.png"))
            button_1 = Button(
                canvas,
                image=GroupChatSendMessageButton,
                borderwidth=0,
                highlightthickness=0,
                command=lambda: on_send_message(),
                relief="flat"
            )
            button_1.place(
                x=277.0,
                y=411.0,
                width=40.0,
                height=36.0
            )

            canvas.create_text(
                12.0,
                0.0,
                anchor="nw",
                text="" +group_name,
                fill="#173054",
                font=("LibreCaslonText Regular", 16 * -1)
            )
            print("Rectangle clicked for group:", group_name)

        # Iterate over documents
        for i, group_chat in enumerate(group_chats):
            group_name = group_chat["GroupName"]  # Get the group name from the document
            
            # Create rectangle
            rectangle_id = canvas.create_rectangle(
                4.0,
                y_offset + 40 * i + distance_between_rectangles * i,  # Adjust y-offset for each group chat
                186.0,
                y_offset + 40 * (i + 1) + distance_between_rectangles * i,  # Adjust y-offset for each group chat
                fill="#3A868F",
                outline="",
                tags=("group_rectangle",)  # Add a tag to identify rectangles
            )
            
            # Bind the click event to the rectangle and pass the group name as an argument
            canvas.tag_bind(rectangle_id, "<Button-1>", lambda event, name=group_name: on_rectangle_click(event, name))
            
            # Create text for group name
            canvas.create_text(
                11.0,
                y_offset + 10 + 40 * i + distance_between_rectangles * i,  # Adjust y-offset for each group chat
                anchor="nw",
                text=group_name,  # Use the group name
                fill="#FFFFFF",
                font=("Libre Caslon Text", 12 * -1)
            )
            
            # Create text for number of selected users
            canvas.create_text(
                169.0,
                y_offset + 20 + 40 * i + distance_between_rectangles * i,  # Adjust y-offset for each group chat
                anchor="nw",
                text=str(len(group_chat["SelectedUsers"])),  # Get number of selected users from document
                fill="#FFFFFF",
                font=("Libre Caslon Text", 12 * -1)
            )
    


    def search_employees(username):
        regex = f"^{username}"
        # Assuming EmployeeCollection is accessible here
        employees = EmployeeCollection.find({"UserName": {"$regex": regex}})
        return [employee["UserName"] for employee in employees]

    def open_group_creation_window():
        # Create a new window for group creation
        group_creation_window = tk.Toplevel()
        group_creation_window.title("Create Group")

        def add_selected_user():
            selected_user = search_dropdown.get()
            if selected_user:
                selected_users_listbox.insert(tk.END, selected_user)
                search_dropdown.set('')
        
        def create_group():
            group_name = group_name_entry.get()
            selected_users = selected_users_listbox.get(0, tk.END)
            
            # Check if group name already exists
            if GroupChatCollection.find_one({"GroupName": group_name}):
                messagebox.showerror("Error", "A group with the same name already exists. Please choose a different name.")
                return
            
            if not group_name:
                messagebox.showerror("Error", "Please enter a group name.")
                return
            if not selected_users:
                messagebox.showerror("Error", "Please select at least one user for the group.")
                return
            
            # Store group chat data in the database
            group_chat_data = {
                "CreatedBy": UserID,  # Assuming UserID is the username of the user who created the group
                "CompanyName": CompanyName,  # Replace with the actual company name
                "GroupName": group_name,
                "SelectedUsers": selected_users
            }
            GroupChatCollection.insert_one(group_chat_data)
            
            print("Group Name:", group_name)
            print("Selected Users:", selected_users)
            created_groups.append(group_name)
            DisplayGroups()
            group_creation_window.destroy()

        def on_search_entry_changed(*args):
            search_text = search_var.get()
            matches = search_employees(search_text)
            search_dropdown.configure(values=matches)

        # Widgets in the group creation window
        group_name_label = tk.Label(group_creation_window, text="Enter Group Name:")
        group_name_label.grid(row=0, column=0, padx=10, pady=5)

        group_name_entry = tk.Entry(group_creation_window)
        group_name_entry.grid(row=0, column=1, padx=10, pady=5)

        search_label = tk.Label(group_creation_window, text="Search Employee:")
        search_label.grid(row=1, column=0, padx=10, pady=5)

        search_var = tk.StringVar()
        search_var.trace("w", on_search_entry_changed)
        search_entry = tk.Entry(group_creation_window, textvariable=search_var)
        search_entry.grid(row=1, column=1, padx=10, pady=5)

        search_dropdown = ttk.Combobox(group_creation_window)
        search_dropdown.grid(row=2, column=0, columnspan=2, padx=10, pady=5)

        add_button = tk.Button(group_creation_window, text="Add", command=add_selected_user)
        add_button.grid(row=3, column=0, columnspan=2, padx=10, pady=5)

        selected_users_label = tk.Label(group_creation_window, text="Selected Users:")
        selected_users_label.grid(row=4, column=0, padx=10, pady=5)

        selected_users_listbox = tk.Listbox(group_creation_window, selectmode="multiple")
        selected_users_listbox.grid(row=5, column=0, columnspan=2, padx=10, pady=5)

        create_group_button = tk.Button(group_creation_window, text="Create Group", command=create_group)
        create_group_button.grid(row=6, column=0, columnspan=2, padx=10, pady=5)

        submit_button = tk.Button(group_creation_window, text="Submit", command=group_creation_window.destroy)
        submit_button.grid(row=7, column=0, columnspan=2, padx=10, pady=5)

    
    # # Main frame
    # main_frame = tk.Frame(window, width=510, height=450)
    # main_frame.place(x=36,y=0)

    # Left frame for group chat label and search button
    left_frame = tk.Frame(GroupChatFrame, width=190, height=450)
    left_frame.place(x=0,y=0)

    canvas = Canvas(
        left_frame,
        bg = "#ECECD9",
        height = 450,
        width = 190,
        bd = 0,
        highlightthickness = 0,
        relief = "ridge"
    )

    canvas.place(x = 0, y = 0)
    
    canvas.create_rectangle(
    0.0,
    0.0,
    190.0,
    39.0,
    fill="#3A868F",
    outline="")

    canvas.create_rectangle(
        4.0,
        43.0,
        186.0,
        81.0,
        fill="#3A868F",
        outline="")

    canvas.create_rectangle(
        4.0,
        87.0,
        186.0,
        125.0,
        fill="#3A868F",
        outline="")

    LeftFramePhotoImageGroupChat = PhotoImage(
        file=relative_to_assets("image_18.png"))
    image_1 = canvas.create_image(
        95.0,
        225.0,
        image=LeftFramePhotoImageGroupChat
    )

    CreateGroupChatButtonImage = PhotoImage(
        file=relative_to_assets("button_37.png"))
    button_1 = Button(
        image=CreateGroupChatButtonImage,
        borderwidth=0,
        highlightthickness=0,
        command=lambda: open_group_creation_window(),
        relief="flat"
    )
    button_1.place(
        x=154.0,
        y=3.0,
        width=33.0,
        height=33.0
    )

    canvas.create_text(
        1.0,
        1.0,
        anchor="nw",
        text="Group Chat",
        fill="#173054",
        font=("Libre Caslon Text", 14 * -1)
    )
    canvas.create_rectangle(
    0.0,
    0.0,
    190.0,
    39.0,
    fill="#3A868F",
    outline="")

    
    # Right frame for group creation
    right_frame = tk.Frame(GroupChatFrame, width=320, height=450)
    right_frame.place(x=190, y=0)
    DisplayGroups()

def AppointmentSchedulingAndDisplay(CompanyName,UserID):
    global BookAnAppointmentButtonImage
    AppointmentSchedulingFrame =  Frame(window,height="450", width="510")
    AppointmentSchedulingFrame.place(x=36,y=0)
    def BookAnAppointment(CompanyName,UserID):
        global entry_bg_1,entry_bg_2,entry_bg_3,BookAnAppointmentSubmitButtonImage
        appointment_window = tk.Toplevel(window)
        appointment_window.title("Book an Appointment")
        appointment_window.geometry("320x450")

        def submit_appointment():
            reason = entry_1.get()
            date = entry_3.get()
            time = entry_2.get()

            # Validate the input (you can add more validation if needed)
            if not reason or not date or not time:
                messagebox.showerror("Error", "Please fill in all fields.")
                return

            # Construct the appointment document
            appointment = {
                "CompanyName":CompanyName,
                "UserName":UserID,
                "reason": reason,
                "date": date,
                "time": time,
                "Status":"Pending"
            }

            # Insert the appointment into the MongoDB collection
            try:
                AppointmentSchedullingCollection.insert_one(appointment)
                messagebox.showinfo("Success", "Appointment scheduled successfully.")
                
            except pymongo.errors.PyMongoError as e:
                messagebox.showerror("Error", f"An error occurred: {e}")

            # Clear the entry fields after submission
            entry_1.delete(0, tk.END)
            entry_2.delete(0, tk.END)
            entry_3.delete(0, tk.END)
            appointment_window.destroy()
            
        
                
        canvas = Canvas(
            appointment_window,
            bg = "#173054",
            height = 450,
            width = 320,
            bd = 0,
            highlightthickness = 0,
            relief = "ridge"
        )

        canvas.place(x = 0, y = 0)
        BookAnAppointmentSubmitButtonImage = PhotoImage(
            file=relative_to_assets("button_50.png"))
        button_1 = Button(
            appointment_window,
            image=BookAnAppointmentSubmitButtonImage,
            borderwidth=0,
            highlightthickness=0,
            command=lambda: submit_appointment(),
            relief="flat"
        )
        button_1.place(
            x=95.0,
            y=322.0,
            width=130.0,
            height=45.0
        )

        canvas.create_text(
            46.0,
            0.0,
            anchor="nw",
            text="Book an Appointment",
            fill="#FFFFFF",
            font=("LibreCaslonText Regular", 20 * -1)
        )

        canvas.create_text(
            33.0,
            95.0,
            anchor="nw",
            text="Reason:",
            fill="#FFFFFF",
            font=("LibreCaslonText Regular", 17 * -1)
        )

        entry_image_1 = PhotoImage(
            file=relative_to_assets("entry_58.png"))
        entry_bg_1 = canvas.create_image(
            160.0,
            129.5,
            image=entry_image_1
        )
        entry_1 = Entry(
            appointment_window,
            bd=0,
            bg="#D9D9D9",
            fg="#000716",
            highlightthickness=0
        )
        entry_1.place(
            x=45.0,
            y=116.0,
            width=230.0,
            height=25.0
        )

        canvas.create_text(
            33.0,
            219.0,
            anchor="nw",
            text="Time:(HH:MM)",
            fill="#FFFFFF",
            font=("LibreCaslonText Regular", 17 * -1)
        )

        entry_image_2 = PhotoImage(
            file=relative_to_assets("entry_59.png"))
        entry_bg_2 = canvas.create_image(
            160.0,
            191.5,
            image=entry_image_2
        )
        entry_2 = Entry(
            appointment_window,
            bd=0,
            bg="#D9D9D9",
            fg="#000716",
            highlightthickness=0
        )
        entry_2.place(
            x=45.0,
            y=178.0,
            width=230.0,
            height=25.0
        )

        canvas.create_text(
            33.0,
            157.0,
            anchor="nw",
            text="Date:(YYYY-MM-DD)",
            fill="#FFFFFF",
            font=("LibreCaslonText Regular", 17 * -1)
        )

        entry_image_3 = PhotoImage(
            file=relative_to_assets("entry_60.png"))
        entry_bg_3 = canvas.create_image(
            160.0,
            253.5,
            image=entry_image_3
        )
        entry_3 = Entry(
            appointment_window,
            bd=0,
            bg="#D9D9D9",
            fg="#000716",
            highlightthickness=0
        )
        entry_3.place(
            x=45.0,
            y=240.0,
            width=230.0,
            height=25.0
        )

    canvas = Canvas(
        AppointmentSchedulingFrame,
        bg = "#173054",
        height = 450,
        width = 510,
        bd = 0,
        highlightthickness = 0,
        relief = "ridge"
    )

    canvas.place(x = 0, y = 0)
    canvas.create_text(
        61.0 + 36,
        7.0,
        anchor="nw",
        text="Appointment Scheduling",
        fill="#FFFFFF",
        font=("LibreCaslonText Regular", 18 * -1)
    )

    BookAnAppointmentButtonImage = PhotoImage(
        file=relative_to_assets("button_49.png"))
    button_1 = Button(
        AppointmentSchedulingFrame,
        image=BookAnAppointmentButtonImage,
        borderwidth=0,
        highlightthickness=0,
        command=lambda: BookAnAppointment(CompanyName,UserID),
        relief="flat"
    )
    button_1.place(
        x=383.0,
        y=13.0,
        width=106.0,
        height=36.0
    )
    

    def display_appointments():

        appointments = AppointmentSchedullingCollection.find({"CompanyName": CompanyName, "UserName": UserID})
        y_offset = 120 
        # Clear previous appointments
        canvas.delete("appointments")
        for appointment in appointments:
            reason = appointment["reason"]
            date = appointment["date"]
            time = appointment["time"]
            status = appointment["Status"]

            # Create a rectangle for each appointment
            canvas.create_rectangle(
                21.0,
                y_offset,
                489.0,
                y_offset + 58.0,  # Adjust height as needed
                fill="#3A868F",
                outline=""
            )

            # Display appointment details
            canvas.create_text(
                288.0,
                y_offset + 23.0,
                anchor="nw",
                text=f"Status: {status}",
                fill="#FFFFFF",
                font=("LibreCaslonText Regular", 15 * -1)
            )

            canvas.create_text(
                30.0,
                y_offset + 6.0,
                anchor="nw",
                text=f"Reason: {reason}",
                fill="#FFFFFF",
                font=("LibreCaslonText Regular", 17 * -1)
            )

            canvas.create_text(
                30.0,
                y_offset + 33.0,
                anchor="nw",
                text=f"Date: {date}",
                fill="#FFFFFF",
                font=("LibreCaslonText Regular", 15 * -1)
            )

            canvas.create_text(
                140.0,
                y_offset + 33.0,
                anchor="nw",
                text=f"Time: {time}",
                fill="#FFFFFF",
                font=("LibreCaslonText Regular", 15 * -1)
            )

            y_offset += 70  # Increment Y offset for the next appointment
    
    def update_appointments_thread():
        while True:
            display_appointments()
            # Update appointments every 15 seconds
            time.sleep(15)

    update_thread = threading.Thread(target=update_appointments_thread)
    update_thread.daemon = True  # Daemonize the thread so it automatically exits when the main program exits
    update_thread.start()
    
def UploadDocument(CompanyName,UserID):
    global UploadDocumentButtonImage
    UploadDocumentFrame = Frame(window, height="450", width="510")
    UploadDocumentFrame.place(x=36,y=0)

    def upload():
        file_path = filedialog.askopenfilename()
        if file_path:
            with open(file_path, 'rb') as file:
                document_data = file.read()
                document_name = os.path.basename(file_path)
                document = {
                    'name': document_name, 
                    'data': Binary(document_data), 
                    'CompanyName': CompanyName, 
                    'shared_with': [UserID]  # Store UserID in the shared_with array
                }
                DocumentCollection.insert_one(document)
                print("Document uploaded successfully.")
    
    def share_file(file_id, selected_users):
        # Update the document in the database to include the new users in shared_with
        DocumentCollection.update_one({"_id": file_id}, {"$push": {"shared_with": {"$each": selected_users}}})
        print(f"File shared with {selected_users} successfully.")
        share_window.destroy()  # Close the share window after submitting

    
    def open_share_window(file_id):
        global share_window,TextBoxtoSelectUsersBGImage,SharetheDocumentButtonImage
        share_window = Toplevel()
        share_window.title("Share File")
        share_window.geometry("437x299")
        share_window.configure(bg = "#173054")
                
        canvas = Canvas(
            share_window,
            bg = "#173054",
            height = 299,
            width = 437,
            bd = 0,
            highlightthickness = 0,
            relief = "ridge"
        )

        canvas.place(x = 0, y = 0)
        canvas.create_text(
            109.0,
            7.0,
            anchor="nw",
            text="Share file",
            fill="#FFFFFF",
            font=("LibreCaslonText Regular", 20 * -1)
        )

        canvas.create_text(
            21.0,
            53.0,
            anchor="nw",
            text="Select User:",
            fill="#FFFFFF",
            font=("LibreCaslonText Regular", 16 * -1)
        )

        # TextBoxtoSelectUsersBGImage = PhotoImage(
        #     file=relative_to_assets("entry_61.png"))
        # entry_bg_1 = canvas.create_image(
        #     269.5,
        #     63.0,
        #     image=TextBoxtoSelectUsersBGImage
        # )
        # entry_1 = Entry(
        #     bd=0,
        #     bg="#D9D9D9",
        #     fg="#000716",
        #     highlightthickness=0
        # )
        # entry_1.place(
        #     x=136.0,
        #     y=50.0,
        #     width=267.0,
        #     height=24.0
        # )

        canvas.create_rectangle(
            100.0,
            95.0,
            336.0,
            203.0,
            fill="#D9D9D9",
            outline="")

        SharetheDocumentButtonImage = PhotoImage(
            file=relative_to_assets("button_52.png"))
        button_1 = Button(
            share_window,
            image=SharetheDocumentButtonImage,
            borderwidth=0,
            highlightthickness=0,
            command=lambda: print("button_1 clicked"),
            relief="flat"
        )
        button_1.place(
            x=161.0,
            y=237.0,
            width=115.0,
            height=43.0
        )
        
        usernames = EmployeeCollection.distinct('UserName', {'CompanyName': CompanyName})
        print(usernames)

        # Function to populate the dropdown menu with usernames
        def populate_dropdown(usernames):
            user_dropdown['values'] = usernames

        # Dropdown menu for selecting users
        user_dropdown = ttk.Combobox(share_window)
        user_dropdown.place(x=136.0, y=50.0, width=267.0, height=24.0)
        populate_dropdown(usernames)  # Populate the dropdown with usernames

        # Listbox for displaying selected users
        selected_users_listbox = tk.Listbox(share_window)
        selected_users_listbox.place(x=100.0, y=95.0, width=236.0, height=108.0)

        
        # Button to add selected user to the listbox
        add_button = tk.Button(share_window, text="Add", command=lambda: selected_users_listbox.insert(tk.END, user_dropdown.get()))
        add_button.place(x=350, y=90.0, width=60.0, height=24.0)

        # Button to submit selected users
        submit_button = Button(share_window, text="Submit", command=lambda: share_file(file_id, selected_users_listbox.get(0, tk.END)))
        submit_button.place(x=161.0, y=237.0, width=115.0, height=43.0)

    def download_file(file_name):
        document = DocumentCollection.find_one({'name': file_name})
        if document:
            file_data = document['data']
            file_path = filedialog.asksaveasfilename(defaultextension="." + os.path.splitext(file_name)[1])
            with open(file_path, 'wb') as file:
                file.write(file_data)
            print("File downloaded successfully.")
                
    canvas = Canvas(
        UploadDocumentFrame,
        bg = "#173054",
        height = 450,
        width = 510,
        bd = 0,
        highlightthickness = 0,
        relief = "ridge"
    )

    canvas.place(x = 0, y = 0)
    canvas.create_text(
        21.0,
        7.0,
        anchor="nw",
        text="Files",
        fill="#FFFFFF",
        font=("LibreCaslonText Regular", 18 * -1)
    )

    UploadDocumentButtonImage = PhotoImage(
        file=relative_to_assets("button_51.png"))
    button_1 = Button(
        UploadDocumentFrame,
        image=UploadDocumentButtonImage,
        borderwidth=0,
        highlightthickness=0,
        command=lambda: upload(),
        relief="flat"
    )
    button_1.place(
        x=383.0,
        y=13.0,
        width=106.0,
        height=36.0
    )


    files = DocumentCollection.find({'CompanyName': CompanyName, 'shared_with': UserID})

    y_offset = 100
    for file in files:
        file_id = file['_id']
        file_name = file['name']
        
        canvas.create_rectangle(
            21.0,
            y_offset,
            489.0,
            y_offset + 31.0,
            fill="#3A868F",
            outline=""
        )
        canvas.create_text(
            30.0,
            y_offset + 8.0,
            anchor="nw",
            text=file_name,
            fill="#FFFFFF",
            font=("LibreCaslonText Regular", 17 * -1)
        )
        
        # Create a button for sharing the file
        share_button = tk.Button(UploadDocumentFrame, text="Share", command=lambda file_id=file_id: open_share_window(file_id))
        share_button.place(x=450, y=y_offset + 5)

        # Create a button for downloading the file
        download_button = tk.Button(UploadDocumentFrame, text="Download", command=lambda file_name=file_name: download_file(file_name))
        download_button.place(x=350, y=y_offset + 5)
        
        y_offset += 40

def UserProfile(CompanyName,UserID):

    global image_image_1,image_image_2,button_image_1,button_image_2,button_image_3,button_image_4,button_image_5,button_image_6
    global button_image_7,button_image_8,button_image_9,button_image_10,button_image_11,button_image_12
    
    user_data = EmployeeCollection.find_one({"UserName": UserID})
    if user_data and "photo" in user_data:
        photo_data_base64 = user_data["photo"]
        # Decode Base64 data
        photo_data_binary = base64.b64decode(photo_data_base64)
        with open("user_photo.jpg", "wb") as f:
            f.write(photo_data_binary)
        image = Image.open("user_photo.jpg")
        image = image.resize((112, 112), Image.LANCZOS)
        image = round_corners(image, 50)
        
    else:
        print("Photo not found for the specified user.")
        return None
    if UserID:
        
        canvas = Canvas(
        window,
        bg = "#173054",
        height = 450,
        width = 700,
        bd = 0,
        highlightthickness = 0,
        relief = "ridge"
    )

        canvas.place(x = 0, y = 0)
        image_image_1 = PhotoImage(
            file=relative_to_assets("image_10.png"))
        image_1 = canvas.create_image(
            350.0,
            225.0,
            image=image_image_1
        )

        canvas.create_rectangle(
            546.0,
            0.0,
            700.0,
            450.0,
            fill="#225777",
            outline="")

        canvas.create_text(
            575.0,
            150.0,
            anchor="nw",
            text=""+ user_data["UserName"],
            fill="#FFFFFF",
            font=("Libre Caslon Text", 16 * -1)
        )

        canvas.create_text(
            557.0,
            229.0,
            anchor="nw",
            text="Tasks",
            fill="#FFFFFF",
            font=("Libre Caslon Text", 12 * -1)
        )

        canvas.create_text(
            575.0,
            171.0,
            anchor="nw",
            text="" + user_data["Designation"],
            fill="#FFFFFF",
            font=("Libre Caslon Text", 8 * -1)
        )

        canvas.create_text(
            557.0,
            343.0,
            anchor="nw",
            text="Todo List",
            fill="#FFFFFF",
            font=("Libre Caslon Text", 12 * -1)
        )

        image_image_2 = PhotoImage(
            file=relative_to_assets("image_11.png"))
        image_2 = canvas.create_image(
            625.0,
            79.0,
            image=image_image_2
        )

        canvas.create_rectangle(
            0.0,
            0.0,
            36.0,
            450.0,
            fill="#225777",
            outline="")

        button_image_1 = PhotoImage(
            file=relative_to_assets("button_13.png"))
        VideoCallButton = Button(
            image=button_image_1,
            borderwidth=0,
            highlightthickness=0,
            command=lambda: print("button_13 clicked"),
            relief="flat"
        )
        VideoCallButton.place(
            x=2.0,
            y=3.0,
            width=32.0,
            height=32.0
        )

        button_image_2 = PhotoImage(
            file=relative_to_assets("button_14.png"))
        MailButton = Button(
            image=button_image_2,
            borderwidth=0,
            highlightthickness=0,
            command=lambda: MailApplication(),
            relief="flat"
        )
        MailButton.place(
            x=2.0,
            y=35.0,
            width=32.0,
            height=32.0
        )

        button_image_3 = PhotoImage(
            file=relative_to_assets("button_15.png"))
        button_3 = Button(
            image=button_image_3,
            borderwidth=0,
            highlightthickness=0,
            command=lambda:Display_current_user_for_VideoCall(UserID),
            relief="flat"
        )
        button_3.place(
            x=2.0,
            y=67.0,
            width=32.0,
            height=32.0
        )

        button_image_4 = PhotoImage(
            file=relative_to_assets("button_16.png"))
        button_4 = Button(
            image=button_image_4,
            borderwidth=0,
            highlightthickness=0,
            command=lambda: get_current_ip_address(UserID),
            relief="flat"
        )
        button_4.place(
            x=2.0,
            y=99.0,
            width=32.0,
            height=32.0
        )

        button_image_5 = PhotoImage(
            file=relative_to_assets("button_17.png"))
        button_5 = Button(
            image=button_image_5,
            borderwidth=0,
            highlightthickness=0,
            command=lambda: GroupChatApplication(CompanyName,UserID),
            relief="flat"
        )
        button_5.place(
            x=2.0,
            y=131.0,
            width=32.0,
            height=32.0
        )

        button_image_6 = PhotoImage(
            file=relative_to_assets("button_18.png"))
        button_6 = Button(
            image=button_image_6,
            borderwidth=0,
            highlightthickness=0,
            command=lambda: print("button_18 clicked"),
            relief="flat"
        )
        button_6.place(
            x=2.0,
            y=165.0,
            width=32.0,
            height=32.0
        )

        button_image_7 = PhotoImage(
            file=relative_to_assets("button_19.png"))
        button_7 = Button(
            image=button_image_7,
            borderwidth=0,
            highlightthickness=0,
            command=lambda: Todo(UserID,CompanyName),
            relief="flat"
        )
        button_7.place(
            x=2.0,
            y=197.0,
            width=32.0,
            height=32.0
        )

        button_image_8 = PhotoImage(
            file=relative_to_assets("button_20.png"))
        button_8 = Button(
            image=button_image_8,
            borderwidth=0,
            highlightthickness=0,
            command=lambda: LeaveManagement(CompanyName,UserID),
            relief="flat"
        )
        button_8.place(
            x=2.0,
            y=230.0,
            width=32.0,
            height=32.0
        )

        button_image_9 = PhotoImage(
            file=relative_to_assets("button_21.png"))
        button_9 = Button(
            image=button_image_9,
            borderwidth=0,
            highlightthickness=0,
            command=lambda: UploadDocument(CompanyName,UserID),
            relief="flat"
        )
        button_9.place(
            x=2.0,
            y=263.0,
            width=32.0,
            height=32.0
        )

        button_image_10 = PhotoImage(
            file=relative_to_assets("button_22.png"))
        button_10 = Button(
            image=button_image_10,
            borderwidth=0,
            highlightthickness=0,
            command=lambda: AppointmentSchedulingAndDisplay(CompanyName,UserID),#appointment
            relief="flat"
        )
        button_10.place(
            x=2.0,
            y=298.0,
            width=32.0,
            height=32.0
        )

        button_image_11 = PhotoImage(
            file=relative_to_assets("button_23.png"))
        button_11 = Button(
            image=button_image_11,
            borderwidth=0,
            highlightthickness=0,
            command=lambda:logout(),
            relief="flat"
        )
        button_11.place(
            x=2.0,
            y=412.0,
            width=32.0,
            height=32.0
        )

        canvas.create_rectangle(
            557.0,
            248.0,
            688.0,
            330.0,
            fill="#173054",
            outline="")

        canvas.create_rectangle(
            557.0,
            362.0,
            688.0,
            444.0,
            fill="#173054",
            outline="")

        button_image_12 = PhotoImage(
            file=relative_to_assets("button_24.png"))
        button_12 = Button(
            image=button_image_12,
            borderwidth=0,
            highlightthickness=0,
            command=lambda: print("button_24 clicked"),
            relief="flat"
        )
        button_12.place(
            x=570.0,
            y=187.0,
            width=106.0,
            height=20.0
        )
        
        if image:
            photo = ImageTk.PhotoImage(image)
            canvas.create_image(
                569.0,
                23.0,
                anchor="nw",
                image=photo,
                )
            canvas.image = photo 
        canvas.create_text(
        569.0,
        20.0,
        anchor="nw",
        text="",
        fill="#FFFFFF",
        font=("Inter", 12 * -1)
        )

        canvas.create_text(
            557.0,
            364.0,
            anchor="nw",
            text="task 1\n\n",
            fill="#FFFFFF",
            font=("LibreCaslonText Regular", 12 * -1)
        )

        canvas.create_text(
            557.0,
            379.0,
            anchor="nw",
            text="task 1\n\n",
            fill="#FFFFFF",
            font=("LibreCaslonText Regular", 12 * -1)
        )

        canvas.create_text(
            557.0,
            395.0,
            anchor="nw",
            text="task 1\n\n",
            fill="#FFFFFF",
            font=("LibreCaslonText Regular", 12 * -1)
        )

        canvas.create_text(
            557.0,
            411.0,
            anchor="nw",
            text="task 1\n\n",
            fill="#FFFFFF",
            font=("LibreCaslonText Regular", 12 * -1)
        )

        canvas.create_text(
            557.0,
            426.0,
            anchor="nw",
            text="task 1\n\n",
            fill="#FFFFFF",
            font=("LibreCaslonText Regular", 12 * -1)
        )
        ip_address = socket.gethostbyname(socket.gethostname())
    
        # Check if the username already exists in the collection
        existing_user = EmployeeCollection.find_one({"UserName": UserID})
        if existing_user:
            # Update the existing document with the new IP address
            EmployeeCollection.update_one({"UserName": UserID}, {"$set": {"ip_address": ip_address}})

        def receive_notifications():
            try:
                document = EmployeeCollection.find_one({})
                SERVER_IP = document.get("NotificationServerIP")
                SERVER_PORT = document.get("NotificationServerPort")
                static_message = f"Join the Meeting. You are invited by {UserID}"
                client_socket = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
                client_socket.connect((SERVER_IP, SERVER_PORT))
                while True:
                    message = client_socket.recv(1024).decode()
                    if message:
                        if message.startswith(static_message):
                            user_invited_by = message.split("by ", 1)[1]
                            print(f"Invited by: {user_invited_by}")
                        messagebox.showinfo("Notification", f"Received: {message}")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to receive notification: {e}")

        # Start receiving notifications in a separate thread
        receive_thread = threading.Thread(target=receive_notifications)
        receive_thread.start()

def JoinMeetingClient(UserName):
    local_ip_address = socket.gethostbyname(socket.gethostname())

    server = StreamingServer(local_ip_address, 7777)
    receiver = AudioReceiver(local_ip_address, 6666)

    def start_listening():
        t1= threading.Thread(target=server.start_server)
        t2= threading.Thread(target=receiver.start_server,)
        t1.start()
        t2.start()


    def start_camera_stream():
        camera_client = CameraClient(text_target_ip,9999)
        t3 = threading.Thread(target=camera_client.start_stream)
        t3.start()

    def start_screen_sharing():
        screen_client = ScreenShareClient(text_target_ip,9999)
        t4 = threading.Thread(target=screen_client.start_stream)
        t4.start()

    def start_audio_stream():
        audio_sender = AudioSender(text_target_ip,8888)
        t5 = threading.Thread(target=audio_sender.start_stream)
        t5.start()
    #GUI

    window = tk.Tk()
    window.title('video conference')
    window.geometry('300x300')

    label_target_ip = tk.Label(window, text="target IP:")
    label_target_ip.pack()

    employee = EmployeeCollection.find_one({'UserName': UserName})
    text_target_ip = employee.get("ip_address")

    btn_listen = tk.Button(window,text="Start listening", width=50, command=start_listening)
    btn_listen.pack(anchor=tk.CENTER,expand=True)

    btn_camera = tk.Button(window,text="Start Camera stream", width=50, command=start_camera_stream)
    btn_camera.pack(anchor=tk.CENTER,expand=True)

    btn_screen = tk.Button(window,text="Start screen sharing", width=50, command=start_screen_sharing)
    btn_screen.pack(anchor=tk.CENTER,expand=True)

    btn_audio = tk.Button(window,text="Start Audio Stream", width=50, command=start_audio_stream)
    btn_audio.pack(anchor=tk.CENTER,expand=True)
    window.mainloop()

def receive_notifications():
    try:
        document = EmployeeCollection.find_one({})
        SERVER_IP = document.get("NotificationServerIP")
        SERVER_PORT = document.get("NotificationServerPort")
        client_socket = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        client_socket.connect((SERVER_IP, SERVER_PORT))
        while True:
            message = client_socket.recv(1024).decode()
            if message:
                user_invited_by = message.split("by ", 1)[1]
                print(f"Invited by: {user_invited_by}")
                messagebox.showinfo("Notification", f"Received: {message}")
                print(user_invited_by)
                JoinMeetingClient(user_invited_by)
    except Exception as e:
        messagebox.showerror("Error", f"Failed to receive notification: {e}")
# Start receiving notifications in a separate thread

def check_username(UserName):
    existing_user = EmployeeCollection.find_one({"UserName": UserName})
    return existing_user is not None

def SignUpAuth(entry_1,entry_2,entry_3,entry_4):
    CompanyName = entry_1.get()
    username = entry_2.get()
    password = entry_3.get()
    confirm_password = entry_4.get()
    
    if check_username(username):
        messagebox.showerror("Error", "Username already exists!")
        return

    if password == confirm_password:
        messagebox.showinfo('Successful Sign Up ','Done')     
        PersonalDetailForm(CompanyName,username,password)
    else:
        messagebox.showerror('Invalid',"Passwords don't match")

def SignUp():
    global image_image_1,entry_image_1,entry_image_2,entry_image_3,entry_image_4,button_image_1, entry_2, entry_3
    
    canvas = Canvas(
    window,
    bg = "#3A868F",
    height = 450,
    width = 700,
    bd = 0,
    highlightthickness = 0,
    relief = "ridge"
    )

    canvas.place(x = 0, y = 0)

    image_image_1 = PhotoImage(
        file=relative_to_assets("image_5.png"))
    image_1 = canvas.create_image(
        350.0,
        225.0,
        image=image_image_1
    )

    canvas.create_text(
        111.0,
        139.0,
        anchor="nw",
        text="Select your Company",
        fill="#FFFFFF",
        font=("Libre Caslon Text", 11 * -1)
    )

    canvas.create_text(
        111.0,
        189.0,
        anchor="nw",
        text="Create Username",
        fill="#FFFFFF",
        font=("Libre Caslon Text", 11 * -1)
    )

    canvas.create_text(
        111.0,
        236.0,
        anchor="nw",
        text="Password",
        fill="#FFFFFF",
        font=("Libre Caslon Text", 11 * -1)
    )

    canvas.create_text(
        111.0,
        284.0,
        anchor="nw",
        text="Confirm Password",
        fill="#FFFFFF",
        font=("Libre Caslon Text", 11 * -1)
    )

    canvas.create_text(
        156.0,
        90.0,
        anchor="nw",
        text="Sign Up",
        fill="#FFFFFF",
        font=("Libre Caslon Text", 23 * -1)
    )

    entry_image_1 = PhotoImage(
        file=relative_to_assets("entry_3.png"))
    entry_bg_1 = canvas.create_image(
        201.0,
        171.0,
        image=entry_image_1
    )
    
    allrecord= [doc["CompanyName"] for doc in AdminCollection.find({},{'_id':0,'CompanyName':1})]
    
    CompanyNames=[]
    for prv in allrecord:
        CompanyNames.append(prv)

    entry_1 = ttk.Combobox(
        canvas,
        values=CompanyNames
    )
    entry_1.place(
        x=102.0,
        y=159.0,
        width=198.0,
        height=22.0
    )
    

    entry_image_2 = PhotoImage(
        file=relative_to_assets("entry_4.png"))
    entry_bg_2 = canvas.create_image(
        201.5,
        217.0,
        image=entry_image_2
    )
    entry_2 = Entry(
        bd=0,
        bg="#FFFFFF",
        fg="#000716",
        highlightthickness=0
    )
    entry_2.place(
        x=103.0,
        y=205.0,
        width=197.0,
        height=22.0
    )
    entry_2.config(background="#173054")
    entry_2.config(highlightthickness=0)

    entry_image_3 = PhotoImage(
        file=relative_to_assets("entry_5.png"))
    entry_bg_3 = canvas.create_image(
        201.5,
        264.0,
        image=entry_image_3
    )
    entry_3 = Entry(
        bd=0,
        bg="#FFFFFF",
        fg="#000716",
        show='*',
        highlightthickness=0
    )
    entry_3.place(
        x=103.0,
        y=252.0,
        width=197.0,
        height=22.0
    )
    entry_3.config(background="#173054")
    entry_3.config(highlightthickness=0)

    entry_image_4 = PhotoImage(
        file=relative_to_assets("entry_6.png"))
    entry_bg_4 = canvas.create_image(
        201.5,
        312.0,
        image=entry_image_4
    )
    entry_4 = Entry(
        bd=0,
        bg="#FFFFFF",
        fg="#000716",
        show='*',
        highlightthickness=0
    )
    entry_4.place(
        x=103.0,
        y=300.0,
        width=197.0,
        height=22.0
    )
    entry_4.config(background="#173054")
    entry_4.config(highlightthickness=0)

    button_image_1 = PhotoImage(
        file=relative_to_assets("button_6.png"))
    button_1 = Button(
        image=button_image_1,
        borderwidth=0,
        highlightthickness=0,
        command=lambda :SignUpAuth(entry_1,entry_2,entry_3,entry_4),
        relief="flat"
    )
    button_1.place(
        x=102.0,
        y=340.0,
        width=198.0,
        height=31.0
    )
    
def Login(UserName,passcode):
    username =  UserName.get()
    password = passcode.get()
    print(username)
    print(password)


    user = EmployeeCollection.find_one({'UserName': username})

    if user and user['Password'] == password:
        print("Login successful")
        messagebox.showinfo('Login Successful', 'Welcome')
        Company_Name = user.get('CompanyName')  # Extracting CompanyName from the document
        UserProfile(Company_Name,username)

    else:
        print("Login failed")
        messagebox.showerror('Login Failed','Invalid Input')
            
def SignInToSkive():
    global image_image_1,image_image_2,button_image_1,button_image_2,entry_image_1,entry_image_2
    canvas = Canvas(
    window,
    bg = "#ECECD9",
    height = 450,
    width = 700,
    bd = 0,
    highlightthickness = 0,
    relief = "ridge"
    )

    canvas.place(x = 0, y = 0)
    image_image_1 = PhotoImage(
        file=relative_to_assets("image_3.png"))
    image_1 = canvas.create_image(
        350.0,
        225.0,
        image=image_image_1
    )

    image_image_2 = PhotoImage(
        file=relative_to_assets("image_4.png"))
    image_2 = canvas.create_image(
        350.0,
        67.0,
        image=image_image_2
    )

    entry_image_1 = PhotoImage(
        file=relative_to_assets("entry_1.png"))
    entry_bg_1 = canvas.create_image(
        350.0,
        142.5,
        image=entry_image_1
    )
    entry_1 = Entry(
        bd=0,
        bg="#FFFFFF",
        fg="#000716",
        font=('Libre Coslon Text',8),
        highlightthickness=0
    )
    entry_1.place(
        x=280.0,
        y=132.0,
        width=140.0,
        height=19.0
    )

    entry_image_2 = PhotoImage(
        file=relative_to_assets("entry_2.png"))
    entry_bg_2 = canvas.create_image(
        350.0,
        185.5,
        image=entry_image_2
    )
    entry_2 = Entry(
        bd=0,
        bg="#FFFFFF",
        fg="#000716",
        font=('Libre Coslon Text',8),
        show='*',
        highlightthickness=0
    )
    entry_2.place(
        x=280.0,
        y=175.0,
        width=140.0,
        height=19.0
    )

    canvas.create_text(
        280.0,
        119.0,
        anchor="nw",
        text="Username",
        fill="#225777",
        font=("Libre Caslon Text", 10 * -1)
    )

    canvas.create_text(
        280.0,
        162.0,
        anchor="nw",
        text="Password",
        fill="#225777",
        font=("Libre Caslon Text", 10 * -1)
    )

    canvas.create_text(
        322.0,
        96.0,
        anchor="nw",
        text="Sign In",
        fill="#225777",
        font=("Libre Caslon Text", 14 * -1)
    )

    button_image_1 = PhotoImage(
        file=relative_to_assets("button_4.png"))
    button_1 = Button(
        image=button_image_1,
        borderwidth=0,
        highlightthickness=0,
        command=lambda: print("button_1 clicked"),
        relief="flat"
    )
    button_1.place(
        x=381.0,
        y=198.0,
        width=40.0,
        height=12.0
    )

    button_image_2 = PhotoImage(
        file=relative_to_assets("button_5.png"))
    button_2 = Button(
        image=button_image_2,
        borderwidth=0,
        highlightthickness=0,
        command= lambda:Login(entry_1,entry_2) ,
        #
        relief="flat"
    )
    button_2.place(
        x=280.0,
        y=221.0,
        width=140.0,
        height=21.0
    )
    
def Main_Window():
    
    canvas = Canvas(
        window,
        bg = "#FFFFFF",
        height = 450,
        width = 700,
        bd = 0,
        highlightthickness = 0,
        relief = "ridge"
    )

    canvas.place(x = 0, y = 0)
    image_image_1 = PhotoImage(
        file=relative_to_assets("image_1.png"))
    image_1 = canvas.create_image(
        350.0,
        225.0,
        image=image_image_1
    )

    image_image_2 = PhotoImage(
        file=relative_to_assets("image_2.png"))
    image_2 = canvas.create_image(
        512.0,
        126.0,
        image=image_image_2
    )

    canvas.create_text(
        437.0,
        174.0,
        anchor="nw",
        text="Skive brings the \nteam together \nwherever you are",
        fill="#FFFFFF",
        font=("Libre Caslon Text", 24 * -1)
    )

    button_image_1 = PhotoImage(
        file=relative_to_assets("button_1.png"))
    button_1 = Button(
        image=button_image_1,
        borderwidth=0,
        highlightthickness=0,
        command=SignInToSkive,
        relief="flat"
    )
    button_1.place(
        x=475.0,
        y=299.0,
        width=134.0,
        height=23.0
    )

    canvas.create_text(
        437.0,
        388.0,
        anchor="nw",
        text="Is your team new to Skive? Create a new",
        fill="#FFFFFF",
        font=("Libre Caslon Text", 9 * -1)
    )

    canvas.create_text(
        479.0,
        326.0,
        anchor="nw",
        text="New to Skive? ",
        fill="#FFFFFF",
        font=("Libre Caslon Text", 8 * -1)
    )

    button_image_2 = PhotoImage(
        file=relative_to_assets("button_2.png"))
    button_2 = Button(
        image=button_image_2,
        borderwidth=0,
        highlightthickness=0,
        command=SignUp,
        relief="flat"
    )
    button_2.place(
        x=540.0,
        y=326.0,
        width=53.0,
        height=12.0
    )

    button_image_3 = PhotoImage(
        file=relative_to_assets("button_3.png"))
    button_3 = Button(
        image=button_image_3,
        borderwidth=0,
        highlightthickness=0,
        command=Workspace,
        relief="flat"
    )
    button_3.place(
        x=621.0,
        y=388.0,
        width=54.0,
        height=12.0
    )
    window.resizable(True,True)
    #window.state("zoomed")
    start_servers()
    
    window.mainloop()

def main():
    
    Main_Window()

if __name__ == "__main__":
    main()

receive_thread = threading.Thread(target=receive_notifications)
receive_thread.start()